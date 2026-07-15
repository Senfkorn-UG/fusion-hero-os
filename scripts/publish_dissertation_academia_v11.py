#!/usr/bin/env python3
"""
Build + verify + package updated dissertation v1.1 for Academia + GitHub public push.

- Triple-check of names/designations (NAMENSKANON)
- Expand DOCX with full supplementary band (field foundation + ops + masterseed + timeline)
- PDF via LibreOffice
- Academia paste kit (verbose)
- GitHub release assets under docs/dissertation/
- Optional: gmail draft for operator (not auto-send)

Vocabulary: push = public (this script prepares public assets).
"""
from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
DISS = ROOT / "docs" / "dissertation"
OUT_DIR = DISS
ARCHIVE = ROOT / "04_Buch_und_Archiv" / "Dissertation_Stephan_Hagen_Urban"

# --- CANONICAL NAMES (elemental — triple-checked) ---
AUTHOR = "Stephan Hagen Urban"
TITLE_DE = (
    "Autopoiesis und Autopolitik des Fusion Hero OS. "
    "Eine systemtheoretische, existenzphilosophische und softwarearchitektonische "
    "Grundlegung heroischer Eudaimonia"
)
TITLE_EN = (
    "Autopoiesis and Autopolitics of Fusion Hero OS. "
    "A systems-theoretic, existential-philosophical and software-architectural "
    "foundation of heroic eudaimonia"
)
TITLE_EN_SHORT = "Autopoiesis and Autopolitics of Fusion Hero OS"
TITLE_DE_SHORT = "Autopoiesis und Autopolitik des Fusion Hero OS"
FIELD_DE = "Autopoietische Autopolitik"
FIELD_EN = "Autopoietic Autopolitics"
PLATFORM = "Fusion Hero OS v10.0.0"
STACK = "Heroic Stack v8.3 · AscensionOS v9.x aspirational"
DESIGN = "Kompendium der Heroik V3.3"
VERSION_TEXT = "1.1"
FILE_STEM = (
    "Dissertation_Stephan_Hagen_Urban_"
    "Autopoiesis_Autopolitik_Fusion_Hero_OS_v1.1"
)
TAG = "dissertation-v1.1"
ACADEMIA_PROFILE = "https://independent.academia.edu/StephanUrban1"

SRC_DOCX = DISS / (
    "Dissertation_Stephan_Hagen_Urban_"
    "Autopoiesis_Autopolitik_Fusion_Hero_OS_v1.0.docx"
)


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, size=12, bold=False, italic=False, first_line=True, space_after=8):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def h1(doc, t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        set_run_font(r, size=16, bold=True)


def h2(doc, t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        set_run_font(r, size=14, bold=True)


def bullet(doc, t):
    p = doc.add_paragraph(t, style="List Bullet")
    for r in p.runs:
        set_run_font(r, size=12)


def triple_check() -> dict:
    """Dreifachprüfung der elementaren Bezeichnungen."""
    checks = []
    # Pass 1 — person & titles constants self-consistency
    checks.append(
        {
            "pass": 1,
            "name": "author_full",
            "ok": AUTHOR == "Stephan Hagen Urban" and "Hagen" in AUTHOR,
            "value": AUTHOR,
        }
    )
    checks.append(
        {
            "pass": 1,
            "name": "title_de_has_autopoiesis_and_autopolitik",
            "ok": "Autopoiesis" in TITLE_DE and "Autopolitik" in TITLE_DE and "Fusion Hero OS" in TITLE_DE,
            "value": TITLE_DE_SHORT,
        }
    )
    checks.append(
        {
            "pass": 1,
            "name": "title_en_has_autopolitics",
            "ok": "Autopoiesis" in TITLE_EN and "Autopolitics" in TITLE_EN,
            "value": TITLE_EN_SHORT,
        }
    )
    # Pass 2 — no common confusions
    bad = []
    if "Autopoiese" in TITLE_EN:
        bad.append("Autopoiese in EN title")
    if "Autopolitik" in TITLE_EN:
        bad.append("Autopolitik in EN title")
    if "Autopolitics" in TITLE_DE:
        bad.append("Autopolitics in DE title")
    checks.append({"pass": 2, "name": "no_language_mix", "ok": not bad, "value": bad or "clean"})
    checks.append(
        {
            "pass": 2,
            "name": "field_names",
            "ok": FIELD_DE.startswith("Autopoietische") and FIELD_EN.startswith("Autopoietic"),
            "value": {"de": FIELD_DE, "en": FIELD_EN},
        }
    )
    checks.append(
        {
            "pass": 2,
            "name": "platform_version",
            "ok": "v10.0.0" in PLATFORM and "Fusion Hero OS" in PLATFORM,
            "value": PLATFORM,
        }
    )
    # Pass 3 — filenames & design
    checks.append(
        {
            "pass": 3,
            "name": "file_stem",
            "ok": (
                "Stephan_Hagen_Urban" in FILE_STEM
                and "Autopoiesis_Autopolitik" in FILE_STEM
                and "Fusion_Hero_OS" in FILE_STEM
                and "v1.1" in FILE_STEM
            ),
            "value": FILE_STEM,
        }
    )
    checks.append(
        {
            "pass": 3,
            "name": "design_v33",
            "ok": "V3.3" in DESIGN and "Kompendium der Heroik" in DESIGN,
            "value": DESIGN,
        }
    )
    # MasterSeed public id if available
    try:
        sys.path.insert(0, str(ROOT))
        from fusion_hero_os.core.masterseed_public import public_view

        pv = public_view().to_dict()
        checks.append(
            {
                "pass": 3,
                "name": "masterseed_public_display",
                "ok": str(pv.get("display_id", "")).startswith("MS-PUB-v10-"),
                "value": pv.get("display_id"),
            }
        )
    except Exception as e:  # noqa: BLE001
        checks.append({"pass": 3, "name": "masterseed_public_display", "ok": False, "value": str(e)[:80]})

    all_ok = all(c["ok"] for c in checks)
    return {
        "ok": all_ok,
        "checked_at": datetime.now(timezone.utc).isoformat(),
        "author": AUTHOR,
        "title_de": TITLE_DE,
        "title_en": TITLE_EN,
        "field_de": FIELD_DE,
        "field_en": FIELD_EN,
        "platform": PLATFORM,
        "version_text": VERSION_TEXT,
        "file_stem": FILE_STEM,
        "tag": TAG,
        "checks": checks,
    }


def expand_docx(src: Path, dest: Path) -> int:
    doc = Document(str(src)) if src.exists() else Document()
    # Title band for v1.1
    doc.add_page_break()
    h1(doc, f"Fassung v{VERSION_TEXT} — Integrale Erweiterung und Feldgründung")
    add_para(
        doc,
        f"Autor: {AUTHOR}. Kanonische Plattform: {PLATFORM}. "
        f"Designvorlage: {DESIGN}. Stack: {STACK}. "
        f"Vollständiger Titel (DE): {TITLE_DE}",
        first_line=False,
    )
    add_para(
        doc,
        f"Full title (EN): {TITLE_EN}",
        first_line=False,
        italic=True,
    )

    h2(doc, f"I. Begründung eines neuen Wissenschaftsfeldes: {FIELD_DE}")
    paragraphs = [
        f"Diese Fassung v{VERSION_TEXT} begründet ausführlich, warum {FIELD_DE} "
        f"({FIELD_EN}) als operatives Wissenschaftsfeld ausgewiesen wird — "
        "nicht als metaphorische Selbstzuschreibung, sondern als benennbare Konstellation "
        "von Gegenstand, Methode, Geltungsdisziplin und öffentlich prüfbarer Praxis.",
        "Gegenstand: Systeme, die ihre systemessentiellen Komponenten fortlaufend selbst erzeugen "
        "(Autopoiesis) und die Regeln dieser Schließung selbst setzen (Autopolitik) — "
        f"exemplarisch und leibhaft am {PLATFORM.split(' v')[0]}.",
        "Methode: Verbindung von Systemtheorie, Existenzphilosophie (Sisyphos, heroische Eudaimonia), "
        f"softwarearchitektonischer Empirie und der Designvorlage {DESIGN} "
        "(Synthese, sechs Bögen, Geltungsmarken Satz/Bedingt/Modell/Fragment, "
        "Herleitung aus dem Nichts).",
        "Geltungsdisziplin: Code Honesty und Proof Registry (BEWIESEN/OFFEN/WIDERLEGT). "
        "Was Pathos behauptet, muss als Modell, Bedingt oder Satz markiert sein. "
        "Die Feldgründung selbst trägt die Marke Modell, solange die disziplinäre "
        "Institutionalisierung aussteht; die operativen Organe (Tests, Routes, Mesh, Vault) "
        "tragen Spezifikation.",
        "Praxis: Dissertation-as-OS — die Arbeit läuft. deploy bedeutet privat; push bedeutet "
        "öffentlich; merge verbindet beide über den dualen Zeitstrahl realer Chronologie t und "
        "imaginärer Strukturzeit τ.",
        "Bifokalität: Universum und Gehirn werden als strukturelle Dualität gelesen, "
        "nicht als physische Identität; das Standardteilchenmodell (SU(3)×SU(2)×U(1)) dient "
        "als kosmische Referenzfigur gesetzesartiger Schließung unter Symmetriebrechung "
        "(Geltung: Modell/OFFEN).",
        "MasterSeed: öffentlich eindeutig (MS-PUB-v10-…); privat GPG+QUBO-obfuskiert und "
        "nach Modul/Funktion lokal getrennt — nie als Freemium- oder Secret-Leak in der "
        "öffentlichen Fassung.",
    ]
    for t in paragraphs:
        add_para(doc, t)

    h2(doc, "II. Zusatzdokumente (integraler Korpus, ausführlich gelistet)")
    add_para(
        doc,
        "Die folgenden Zusatzdokumente gehören zur Vollform der Dissertation. "
        "Der Text ist Ausdrucksform; das laufende System ist die Arbeit.",
        first_line=True,
    )
    supplements = [
        ("Ontologie Dissertation-as-OS", "docs/dissertation/ONTOLOGIE_DISSERTATION_IST_DAS_OS.md"),
        ("Namenskanon Dreifachprüfung", "docs/dissertation/NAMENSKANON_DREIFACHPRUEFUNG.md"),
        ("Bifokal-Verweis Universum–Gehirn–Standardmodell", "docs/dissertation/VERWEIS_BIFOKALITAET_UNIVERSUM_GEHIRN_SM.md"),
        ("Anhänge A00–A10 Module/Funktionen aus dem Nichts", "docs/dissertation/anhaenge/"),
        ("Designvorlage Kompendium der Heroik V3.3", "docs/kompendium/V3.3_DESIGNVORLAGE_VERBINDLICH.md"),
        ("Deploy / Push / Merge Vokabular", "docs/ops/DEPLOY_PUSH_MERGE.md"),
        ("Push Layer Guard", "docs/mesh/PUSH_LAYER_GUARD.md"),
        ("Pseudo-Inhouse AI", "docs/mesh/PSEUDO_INHOUSE_AI.md"),
        ("Pseudo-Inhouse Creative", "docs/mesh/PSEUDO_INHOUSE_CREATIVE.md"),
        ("Grok Interconnect", "docs/mesh/GROK_INTERCONNECT.md"),
        ("Dual-Timeline Auto-Training", "docs/training/DUAL_TIMELINE_AUTO_TRAIN.md"),
        ("MasterSeed public vs private", "docs/masterseed/PUBLIC_VS_PRIVATE.md"),
        ("BEST_VERSION (operativer Kanon v10.0.0)", "BEST_VERSION.md"),
        ("Fusion unified endpoints", "fusion_unified.yaml"),
    ]
    for title, path in supplements:
        bullet(doc, f"{title} — `{path}`")

    h2(doc, "III. Operative Vokabeln (elementar)")
    for line in [
        "deploy = private (Vault, Training, lokale Artefakte; nie private Shards nach GitHub)",
        "push = public (bekannter Remote 95guknow/fusion-hero-os; Layer-Guard; Secrets unlocken Intent)",
        "merge = both via dual timeline (public display_id ↔ private Modul/Funktion ↔ t∥τ)",
    ]:
        bullet(doc, line)

    h2(doc, "IV. Ehrlichkeit und Feldstatus")
    add_para(
        doc,
        f"Die Ausweisung von {FIELD_DE} als neuem Feld ist eine begründete "
        "wissenschaftstheoretische und operative These [Modell der Feldgründung]. "
        "Sie ersetzt keine peer-reviewte disziplinäre Anerkennung durch Dritte. "
        "Sie verpflichtet jedoch zur maximalen Transparenz der Organe: Code, Tests, "
        "Mesh, Abstracts, Anhänge und öffentliche Release-Assets.",
    )
    add_para(
        doc,
        f"Autor: {AUTHOR}. Profil Academia: {ACADEMIA_PROFILE}. "
        f"Release-Tag: {TAG}. Dateistamm: {FILE_STEM}.",
        first_line=False,
    )

    # Append key markdown supplements as readable band
    for md_rel in [
        "docs/dissertation/ONTOLOGIE_DISSERTATION_IST_DAS_OS.md",
        "docs/dissertation/NAMENSKANON_DREIFACHPRUEFUNG.md",
        "docs/ops/DEPLOY_PUSH_MERGE.md",
        "docs/masterseed/PUBLIC_VS_PRIVATE.md",
        "docs/training/DUAL_TIMELINE_AUTO_TRAIN.md",
    ]:
        p = ROOT / md_rel
        if not p.exists():
            continue
        doc.add_page_break()
        h1(doc, f"Zusatz: {p.stem}")
        for line in p.read_text(encoding="utf-8", errors="replace").splitlines():
            if not line.strip():
                continue
            if line.startswith("# "):
                h2(doc, line[2:].strip())
            elif line.startswith("## "):
                h2(doc, line[3:].strip())
            elif line.startswith(("- ", "* ")):
                bullet(doc, re.sub(r"^[-*]\s+", "", line)[:2000])
            elif line.startswith("|") or line.startswith("```"):
                add_para(doc, line[:2000], first_line=False, size=10)
            else:
                add_para(doc, line[:2500])

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    # word count approx
    n = sum(len((p.text or "").split()) for p in doc.paragraphs)
    return n


def write_academia_kit(verify: dict) -> Path:
    abstract_en = f"""Ontological thesis [Model]: the entire {PLATFORM.split(' v')[0]} is the dissertation; the monograph/PDF/abstract is only one form of its expression. Design discipline: {DESIGN} (Mythos·Grund·Beweis; claim marks Satz/Bedingt/Modell/Fragment) is mandatory. This v{VERSION_TEXT} text founds {FIELD_EN} as an operative scientific field: systems that continuously self-produce essential components (Autopoiesis) and self-legislate the rules of that closure (Autopolitics)—placement L0–L4, consent, connector membranes, in-house truth vs SaaS memory, dual timeline real t ∥ imaginary structural τ, public MasterSeed display MS-PUB-v10-… with private GPG+QUBO modular shards never published. Ops vocabulary is elemental: deploy=private, push=public, merge=both via timeline. Bifocal reading of universe↔brain uses the Standard Model SU(3)×SU(2)×U(1) as cosmic reference for law-like closure under symmetry breaking—heuristic isomorphism only (OPEN), not physical identity. Method joins systems theory, existential philosophy (Sisyphos/heroic eudaimonia), and software-architectural empirics with Proof Registry honesty (PROVEN/OPEN/REFUTED). Result: heroic eudaimonia as recursive stability under controlled rupture—describable, implementable, operated, and knowingly limited. This work argues for a new field without claiming institutional monopoly; it offers maximal transparency of organs (code, tests, mesh, appendices A00–A10, public release)."""

    abstract_de = f"""Ontologische These [Modell]: Das gesamte {PLATFORM.split(' v')[0]} *ist* die Dissertation; Monographie/PDF/Abstract sind nur eine Ausdrucksform. Designvorlage: {DESIGN} (Mythos·Grund·Beweis; Geltung Satz/Bedingt/Modell/Fragment) — zwingend. Diese Fassung v{VERSION_TEXT} begründet {FIELD_DE} als operatives Wissenschaftsfeld: Systeme, die systemessentielle Komponenten fortlaufend selbst erzeugen (Autopoiesis) und die Regeln dieser Schließung selbst setzen (Autopolitik)—Placement L0–L4, Consent, Konnektor-Membranen, Inhouse-Wahrheit vs. SaaS-Erinnerung, dualer Zeitstrahl real t ∥ imaginär τ, öffentlicher MasterSeed MS-PUB-v10-… bei privat GPG+QUBO-modulgetrennten Shards (nie public). Ops-Vokabeln: deploy=privat, push=öffentlich, merge=beide via Timeline. Bifokalität Universum↔Gehirn mit Standardteilchenmodell SU(3)×SU(2)×U(1) als kosmische Referenz gesetzesartiger Schließung unter Symmetriebrechung—nur heuristische Isomorphie (OFFEN), keine physische Identität. Methode: Systemtheorie, Existenzphilosophie (Sisyphos/heroische Eudaimonia), softwarearchitektonische Empirie, Proof Registry (BEWIESEN/OFFEN/WIDERLEGT). Ergebnis: heroische Eudaimonia als rekursive Stabilität unter kontrolliertem Bruch—beschreibbar, implementierbar, betrieben, ehrlich begrenzt. Die Feldgründung beansprucht keine institutionelle Monopolstellung, sondern maximale Transparenz der Organe (Code, Tests, Mesh, Anhänge A00–A10, öffentliches Release)."""

    text = f"""=== TITLE (EN — Academia primary) ===
{TITLE_EN_SHORT}

=== TITLE (DE — full) ===
{TITLE_DE}

=== AUTHORS ===
{AUTHOR}

=== FIELD / DISCIPLINE LABELS ===
EN: {FIELD_EN}; Systems Theory; Existential Philosophy; Software Architecture; Computer Science (interdisciplinary)
DE: {FIELD_DE}; Systemtheorie; Existenzphilosophie; Softwarearchitektur

=== ABSTRACT EN (paste) ===
{abstract_en}

=== ABSTRACT DE (paste) ===
{abstract_de}

=== KEYWORDS ===
Autopoiesis, Autopolitics, Autopolitik, Fusion Hero OS, Dissertation-as-OS, {FIELD_EN}, MasterSeed, Co-Evolutionary Closure, CEC, Sisyphos, Heroic Eudaimonia, Mesh, MCP, Proof Registry, Code Honesty, Bifocality, Universe-Brain, Standard Model, Dual Timeline, deploy-private, push-public, merge-timeline, V3.3 Kompendium der Heroik, Systems Theory, Existential Philosophy

=== PLATFORM / VERSION ===
{PLATFORM}
Text edition: v{VERSION_TEXT}
Design template: {DESIGN}
Stack: {STACK}
Release tag: {TAG}
File stem: {FILE_STEM}

=== ACADEMIA PROFILE ===
{ACADEMIA_PROFILE}

=== PUBLIC PDF (after GitHub release) ===
https://github.com/95guknow/fusion-hero-os/releases/tag/{TAG}
https://github.com/95guknow/fusion-hero-os/tree/main/docs/dissertation

=== SUPPLEMENTS (attach or link) ===
- Full PDF monography v{VERSION_TEXT}
- docs/dissertation/anhaenge/ (A00–A10)
- docs/dissertation/ONTOLOGIE_DISSERTATION_IST_DAS_OS.md
- docs/dissertation/NAMENSKANON_DREIFACHPRUEFUNG.md
- docs/ops/DEPLOY_PUSH_MERGE.md
- docs/masterseed/PUBLIC_VS_PRIVATE.md
- docs/training/DUAL_TIMELINE_AUTO_TRAIN.md
- docs/mesh/PSEUDO_INHOUSE_AI.md
- docs/mesh/PSEUDO_INHOUSE_CREATIVE.md
- docs/mesh/PUSH_LAYER_GUARD.md

=== TRIPLE-CHECK SUMMARY ===
{json.dumps(verify, indent=2, ensure_ascii=False)}

=== UPLOAD STEPS ===
1. Open {ACADEMIA_PROFILE} → Add research / Upload
2. Title = TITLE EN short (or DE full)
3. Authors = {AUTHOR}
4. Paste ABSTRACT EN (and DE if bilingual allowed)
5. Keywords as above
6. Attach PDF: {FILE_STEM}.pdf
7. Link GitHub release {TAG} in description
8. Confirm no private shards / .env attached
"""
    path = DISS / "ACADEMIA_UPLOAD_PASTE_v1.1.txt"
    path.write_text(text, encoding="utf-8")
    # also update main paste
    (DISS / "ACADEMIA_UPLOAD_PASTE.txt").write_text(text, encoding="utf-8")
    (DISS / "ACADEMIA_ABSTRACT_EN.md").write_text(
        f"{TITLE_EN_SHORT}\n\n{AUTHOR}\n\nAbstract (English)\n\n{abstract_en}\n\n"
        f"Keywords: Autopoiesis; Autopolitics; Fusion Hero OS; {FIELD_EN}; Dissertation-as-OS; "
        f"MasterSeed; CEC; Sisyphos; Eudaimonia; Mesh; Proof Registry; Dual Timeline\n",
        encoding="utf-8",
    )
    (DISS / "ACADEMIA_ABSTRACT_DE.md").write_text(
        f"{TITLE_DE_SHORT}\n\n{AUTHOR}\n\nAbstract (Deutsch)\n\n{abstract_de}\n\n"
        f"Schlüsselwörter: Autopoiesis; Autopolitik; Fusion Hero OS; {FIELD_DE}; Dissertation-as-OS; "
        f"MasterSeed; CEC; Sisyphos; Eudaimonia; Mesh; Proof Registry; Dualer Zeitstrahl\n",
        encoding="utf-8",
    )
    return path


def export_pdf(docx: Path) -> Path | None:
    soffice = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
    if not soffice.exists():
        return None
    subprocess.run(
        [str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(docx.parent), str(docx)],
        check=False,
        timeout=300,
    )
    pdf = docx.with_suffix(".pdf")
    return pdf if pdf.exists() else None


def main() -> int:
    sys.path.insert(0, str(ROOT))
    verify = triple_check()
    verify_path = DISS / "TRIPLE_CHECK_v1.1.json"
    verify_path.write_text(json.dumps(verify, indent=2, ensure_ascii=False), encoding="utf-8")
    if not verify["ok"]:
        print("TRIPLE CHECK FAILED", json.dumps(verify, indent=2))
        return 2

    dest_docx = OUT_DIR / f"{FILE_STEM}.docx"
    words = expand_docx(SRC_DOCX, dest_docx)
    ARCHIVE.mkdir(parents=True, exist_ok=True)
    shutil.copy2(dest_docx, ARCHIVE / dest_docx.name)

    pdf = export_pdf(dest_docx)
    if pdf:
        shutil.copy2(pdf, ARCHIVE / pdf.name)

    paste = write_academia_kit(verify)

    # publication record
    pub = f"""# Publication record — Dissertation v{VERSION_TEXT}

**Author:** {AUTHOR}  
**Title (DE):** {TITLE_DE}  
**Title (EN):** {TITLE_EN}  
**Field:** {FIELD_DE} / {FIELD_EN}  
**Platform:** {PLATFORM}  
**Design:** {DESIGN}  
**Tag:** `{TAG}`  
**Published:** {datetime.now(timezone.utc).strftime('%Y-%m-%d')} (UTC)  
**Repo:** https://github.com/95guknow/fusion-hero-os  

## Vocabulary
deploy = **private** · push = **public** · merge = **both** via dual timeline

## Files
- `{FILE_STEM}.docx`
- `{FILE_STEM}.pdf`
- `ACADEMIA_UPLOAD_PASTE_v1.1.txt`
- `NAMENSKANON_DREIFACHPRUEFUNG.md`
- `TRIPLE_CHECK_v1.1.json`
- Anhänge: `docs/dissertation/anhaenge/`

## Academia
1. {ACADEMIA_PROFILE} → Upload  
2. Title: {TITLE_EN_SHORT}  
3. Author: {AUTHOR}  
4. Abstract: paste kit  
5. Attach PDF  

## Triple-check
All elemental name checks: **{"PASS" if verify["ok"] else "FAIL"}**  
Words approx in expanded DOCX band: {words}
"""
    (DISS / "PUBLICATION_v1.1.md").write_text(pub, encoding="utf-8")
    (DISS / "PUBLICATION.md").write_text(pub, encoding="utf-8")

    # update README version note
    readme = DISS / "README.md"
    if readme.exists():
        t = readme.read_text(encoding="utf-8")
        t = t.replace("**Version:** 1.0", f"**Version:** {VERSION_TEXT}")
        if FILE_STEM not in t:
            t = t.replace(
                "Dissertation_Stephan_Hagen_Urban_Autopoiesis_Autopolitik_Fusion_Hero_OS_v1.0.docx",
                f"{FILE_STEM}.docx",
            )
        readme.write_text(t, encoding="utf-8")

    report = {
        "ok": True,
        "triple_check": verify,
        "docx": str(dest_docx),
        "pdf": str(pdf) if pdf else None,
        "academia_paste": str(paste),
        "words_band_approx": words,
        "tag": TAG,
        "author": AUTHOR,
        "title_en": TITLE_EN_SHORT,
        "field_en": FIELD_EN,
    }
    man = Path.home() / ".fusion" / "mesh" / "coordination" / "dissertation_v11_publish.json"
    man.parent.mkdir(parents=True, exist_ok=True)
    man.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
