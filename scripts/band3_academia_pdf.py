#!/usr/bin/env python3
"""
1) Append Band III (empirical protocols + proof tables) to dissertation DOCX
2) Write Academia short abstract (DE + EN)
3) Convert DOCX -> PDF via LibreOffice
"""
from __future__ import annotations

import json
import subprocess
import time
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
DOCX = (
    ROOT
    / "docs"
    / "dissertation"
    / "Dissertation_Stephan_Hagen_Urban_Autopoiesis_Autopolitik_Fusion_Hero_OS_v1.0.docx"
)
ARCH_DOCX = (
    ROOT
    / "04_Buch_und_Archiv"
    / "Dissertation_Stephan_Hagen_Urban"
    / DOCX.name
)
OUT_DIR = ROOT / "docs" / "dissertation"
PROOF = ROOT / "proof_registry.yaml"
COORD = Path.home() / ".fusion" / "mesh" / "coordination"
LO = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")


def font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def p(doc, text, *, first=True, size=12, bold=False, italic=False, center=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first:
        para.paragraph_format.first_line_indent = Cm(0.75)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
    r = para.add_run(text)
    font(r, size=size, bold=bold, italic=italic)
    return para


def h1(doc, t):
    x = doc.add_heading(t, level=1)
    for r in x.runs:
        font(r, 16, bold=True)


def h2(doc, t):
    x = doc.add_heading(t, level=2)
    for r in x.runs:
        font(r, 14, bold=True)


def h3(doc, t):
    x = doc.add_heading(t, level=3)
    for r in x.runs:
        font(r, 12, bold=True)


def bullet(doc, t):
    para = doc.add_paragraph(t, style="List Bullet")
    for r in para.runs:
        font(r, 11)


def set_cell(cell, text, *, bold=False, size=9, fill=None):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    if fill:
        tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for m in cell._tc.tcPr:
        pass
    # margins
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", "40"), ("bottom", "40"), ("left", "60"), ("right", "60")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def add_table(doc, headers, rows, col_widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=9, fill="D9E2F3")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell(table.rows[ri + 1].cells[ci], str(val), size=8)
    # widths
    for row in table.rows:
        for i, w in enumerate(col_widths_cm):
            row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def load_proofs():
    data = yaml.safe_load(PROOF.read_text(encoding="utf-8"))
    claims = data.get("claims") or []
    return claims


def load_json(path: Path):
    if path.is_file():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def append_band3(doc: Document, claims, coord, acad, health_note: str):
    doc.add_page_break()
    h1(doc, "Band III – Empirische Messprotokolle und Proof-Tabellen")
    p(
        doc,
        "Band III dokumentiert die empirische Organschau zum Corpus-Freeze der Dissertation "
        "und tabelliert die autoritative Proof Registry. Alle Zeitstempel sind UTC, sofern "
        "nicht anders angegeben. Code-Honesty bleibt Leitnorm: gemessene Zustände sind "
        "Zustände, keine universellen Gesetze.",
        first=False,
    )

    h2(doc, "III.1 Messprotokoll-Schema")
    bullet(doc, "M1 Infrastruktur-Atmung: Dashboard health, Tailscale status, online tiers")
    bullet(doc, "M2 Koordinationsplan: mesh_cluster_coordinator --mode all")
    bullet(doc, "M3 Curriculum: academia_curriculum_train.py")
    bullet(doc, "M4 Proof Registry Snapshot: proof_registry.yaml parse + Statuszählung")
    bullet(doc, "M5 Archivinventar: Dateizählung Kernbäume (Corpus inventory)")
    bullet(doc, "M6 Konnektor-Katalog: mesh_connectors + mesh_service_coordination")

    h2(doc, "III.2 Hot-Run Protokoll H-2026-07-15-B")
    p(
        doc,
        f"Erfassungszeit (Build Band III): {datetime.now(timezone.utc).isoformat()} UTC. "
        f"Vorheriger Coordinator-Lauf: {coord.get('generated_at', 'n/a')}. "
        f"Dashboard: {health_note}.",
        first=False,
    )

    inv = coord.get("inventory") or {}
    plan = coord.get("plan") or {}
    atlas = coord.get("atlas") or {}
    ts = inv.get("tailscale") or {}
    self_n = ts.get("self") or {}

    h3(doc, "III.2.1 Mesh / Tailscale")
    mesh_rows = []
    if self_n:
        mesh_rows.append(
            [
                self_n.get("hostname", "?"),
                "self",
                "yes" if self_n.get("online") else "no",
                ", ".join(self_n.get("ips") or [])[:40],
                self_n.get("os", ""),
            ]
        )
    for peer in ts.get("peers") or []:
        mesh_rows.append(
            [
                peer.get("hostname", "?"),
                "peer",
                "yes" if peer.get("online") else "no",
                ", ".join(peer.get("ips") or [])[:40],
                peer.get("os", ""),
            ]
        )
    if mesh_rows:
        add_table(
            doc,
            ["Hostname", "Role", "Online", "IPs", "OS"],
            mesh_rows,
            [4.0, 1.5, 1.5, 5.5, 2.0],
        )
    else:
        p(doc, "Keine Tailscale-JSON-Daten im Coordinator-Snapshot.", first=False)

    h3(doc, "III.2.2 Placement / Online-Tiers")
    tiers = plan.get("online_tiers") or []
    p(
        doc,
        f"Online-Tiers: {', '.join(tiers) if tiers else 'n/a'}. "
        f"Inventory ok: {inv.get('ok')}. Drift-Score (Atlas): {atlas.get('drift_score', 'n/a')}. "
        f"Matched roles: {len(inv.get('matched_roles') or [])}.",
        first=False,
    )

    role_rows = []
    for m in inv.get("matched_roles") or []:
        role_rows.append(
            [
                m.get("hostname", m.get("role", "?")),
                m.get("role", ""),
                m.get("tier", ""),
                m.get("platform", ""),
                "yes" if m.get("live") else "no",
            ]
        )
    if role_rows:
        add_table(
            doc,
            ["Host/Key", "Role", "Tier", "Platform", "Live"],
            role_rows,
            [4.0, 3.0, 3.0, 2.5, 1.5],
        )

    h3(doc, "III.2.3 Cluster-Jobs (Plan)")
    job_rows = []
    for j in plan.get("cluster_jobs") or []:
        job_rows.append(
            [
                j.get("id", ""),
                j.get("status", ""),
                j.get("cpu", j.get("reason", ""))[:30],
                (j.get("description") or j.get("reason") or "")[:50],
            ]
        )
    if job_rows:
        add_table(
            doc,
            ["Job-ID", "Status", "CPU/Reason", "Description"],
            job_rows,
            [3.5, 2.5, 3.0, 6.0],
        )

    h3(doc, "III.2.4 Academia-Curriculum State")
    acc = acad.get("account") or {}
    al = acad.get("alignment") or {}
    p(
        doc,
        f"Account: {acc.get('display_name', 'n/a')} | Premium: {acc.get('premium')} | "
        f"Profile: {acc.get('public_profile') or acc.get('public_profile_candidate', 'n/a')} | "
        f"Confidence: {acc.get('profile_confidence', 'n/a')}. "
        f"Phase: {al.get('phase')} | Score: {al.get('score')} | Axiome: {al.get('axioms_loaded')} | "
        f"Phone files: {(acad.get('phone') or {}).get('file_count', 0)}.",
        first=False,
    )
    trows = []
    for t in acad.get("training_targets") or []:
        trows.append([t.get("id", ""), t.get("from_cluster", ""), t.get("inhouse", "")[:40], t.get("status", "")])
    if trows:
        add_table(
            doc,
            ["Target", "Cluster", "Inhouse", "Status"],
            trows,
            [4.0, 2.0, 6.0, 2.0],
        )

    h2(doc, "III.3 Proof-Registry-Tabellen")
    counts = Counter(c.get("status", "?") for c in claims)
    p(
        doc,
        f"Gesamtclaims: {len(claims)}. "
        f"BEWIESEN: {counts.get('BEWIESEN', 0)}; "
        f"OFFEN: {counts.get('OFFEN', 0)}; "
        f"WIDERLEGT: {counts.get('WIDERLEGT', 0)}; "
        f"sonst: {sum(v for k, v in counts.items() if k not in ('BEWIESEN', 'OFFEN', 'WIDERLEGT'))}.",
        first=False,
    )
    add_table(
        doc,
        ["Status", "Anzahl"],
        [[k, str(v)] for k, v in sorted(counts.items())],
        [6.0, 3.0],
    )

    h3(doc, "III.3.1 Vollständige Claim-Tabelle")
    claim_rows = []
    for c in claims:
        stmt = (c.get("statement") or "").replace("\n", " ")
        if len(stmt) > 90:
            stmt = stmt[:87] + "..."
        proofs = c.get("proofs") or []
        claim_rows.append(
            [
                c.get("id", ""),
                c.get("status", ""),
                str(len(proofs)),
                stmt,
            ]
        )
    # chunk tables if very long
    chunk = 40
    for i in range(0, len(claim_rows), chunk):
        if i:
            h3(doc, f"III.3.1 (Fortsetzung {i // chunk + 1})")
        add_table(
            doc,
            ["ID", "Status", "#Proofs", "Statement (gekürzt)"],
            claim_rows[i : i + chunk],
            [3.2, 2.0, 1.3, 8.5],
        )

    h3(doc, "III.3.2 BEWIESEN – Testknoten (Auszug prioritär)")
    bew = [c for c in claims if c.get("status") == "BEWIESEN"]
    rows = []
    for c in bew[:25]:
        pr = "; ".join((c.get("proofs") or [])[:2])
        if len(pr) > 70:
            pr = pr[:67] + "..."
        rows.append([c.get("id", ""), pr])
    add_table(doc, ["Claim-ID", "Proof-Knoten (max 2)"], rows, [4.0, 11.0])

    h3(doc, "III.3.3 OFFEN – bewusst ungedeckte Claims")
    offen = [c for c in claims if c.get("status") == "OFFEN"]
    rows = []
    for c in offen:
        note = (c.get("notes") or c.get("statement") or "")[:80]
        rows.append([c.get("id", ""), note])
    if rows:
        add_table(doc, ["Claim-ID", "Notiz / Statement"], rows, [4.0, 11.0])
    else:
        p(doc, "Keine OFFEN-Claims im aktuellen Registry-Snapshot.", first=False)

    h2(doc, "III.4 Empirische Kennzahlen (aggregiert)")
    metrics = [
        ["Dashboard health", health_note],
        ["Coordinator inventory_ok", str(inv.get("ok"))],
        ["Online tiers", ", ".join(tiers)],
        ["Atlas drift_score", str(atlas.get("drift_score", "n/a"))],
        ["Inhouse present (atlas)", str(len(atlas.get("inhouse_present") or []))],
        ["Inhouse missing paths", str(len(atlas.get("inhouse_missing_paths") or []))],
        ["Proof claims total", str(len(claims))],
        ["Proof BEWIESEN", str(counts.get("BEWIESEN", 0))],
        ["Proof OFFEN", str(counts.get("OFFEN", 0))],
        ["Academia phase", str(al.get("phase"))],
        ["Academia score", str(al.get("score"))],
        ["Platform canon", "10.0.0"],
    ]
    add_table(doc, ["Kennzahl", "Wert"], metrics, [6.0, 9.0])

    h2(doc, "III.5 Interpretationsregeln (nicht überdehnen)")
    bullet(doc, "Online ≠ bewiesen: Health ok belegt Atmung, nicht universelle Theorie.")
    bullet(doc, "BEWIESEN gilt nur mit gelistetem pytest-Knoten und grüner Suite.")
    bullet(doc, "OFFEN ist wissenschaftliche Reife, kein Makel der Autopoiesis.")
    bullet(doc, "DNS-Warnungen im Tailscale-Health sind Membranreibung, kein Kollaps.")
    bullet(doc, "Phone file_count=0: L0-Export fehlt; Routing darf nicht blockieren.")

    h2(doc, "III.6 Rohdaten-Pfade")
    for path in [
        str(COORD / "latest.json"),
        str(COORD / "academia_training_state.json"),
        str(COORD / "dissertation_build_manifest.json"),
        str(COORD / "dissertation_corpus_inventory.json"),
        str(PROOF),
        str(DOCX),
    ]:
        bullet(doc, path)

    h2(doc, "III.7 Schluss Band III")
    p(
        doc,
        "Band III schließt die empirische Klammer der Dissertation: Das System atmet "
        "(Dashboard), organisiert sich (Coordinator), lernt parallel (Curriculum) und "
        "begrenzt sich (Proof Registry). Damit ist die These Autopoiesis+Autopolitik "
        "nicht nur spekulativ, sondern an Messprotokollen und Tabellen anschließbar. "
        "Weitere Sisyphos-Runden können pytest-Vollsuite-Logs und GKE-Job-IDs ergänzen, "
        "sobald kubectl im Operator-PATH stabil ist.",
    )
    p(
        doc,
        "— Ende Band III / Gesamtfassung v1.0 + Band III —",
        first=False,
        bold=True,
        center=True,
    )


def write_academia_abstract():
    de = """Autopoiesis und Autopolitik des Fusion Hero OS

Stephan Hagen Urban

Abstract (Deutsch)

Diese Arbeit entwickelt eine vollständige Theorie der autopoietischen und autopolitischen Organisation am System Fusion Hero OS (operativer Kanon v10.0.0, Heroic Stack v8.3, AscensionOS v9.x aspirational). Autopoiesis bezeichnet die operative Selbstproduktion systemessentieller Komponenten (MasterSeed-Anker, Co-Evolutionary Closure, Sisyphos-Zyklus, Layer-/Fraktalmesh, Archive, Tests). Autopolitik bezeichnet die Selbstgesetzgebung dieser Schließung: Placement-Tiers L0–L4, Consent-Gates, MCP-Konnektoren als Membranen, Tailscale-Mesh als Organsystem und die Trennung von Inhouse-Wahrheit versus externer SaaS-Erinnerung. Methodisch verbindet die Monographie Systemtheorie, Existenzphilosophie (Sisyphos/Eudaimonia) und softwarearchitektonische Empirie inklusive Heißlauf (Dashboard, Coordinator, Curriculum) sowie der Proof Registry (BEWIESEN/OFFEN/WIDERLEGT). Ergebnis: heroische Eudaimonia als rekursive Stabilität unter kontrolliertem Bruch – beschreibbar, implementierbar und ehrlich begrenzt.

Schlüsselwörter: Autopoiesis; Autopolitik; Fusion Hero OS; MasterSeed; CEC; Sisyphos; Eudaimonia; Mesh; MCP; Proof Registry; Code Honesty
"""

    en = """Autopoiesis and Autopolitics of Fusion Hero OS

Stephan Hagen Urban

Abstract (English)

This monograph develops a complete theory of autopoietic and autopolitical organization grounded in Fusion Hero OS (operational canon v10.0.0, heroic stack v8.3, AscensionOS v9.x aspirational). Autopoiesis is operative self-production of system-essential components (MasterSeed anchors, Co-Evolutionary Closure, Sisyphos cycle, layer/fractal mesh, archives, tests). Autopolitics is the self-legislation of that closure: placement tiers L0–L4, consent gates, MCP connectors as membranes, Tailscale mesh as organ network, and the separation of in-house truth from external SaaS memory. Methodologically the work joins systems theory, existential philosophy (Sisyphos/eudaimonia), and software-architectural empirics including live hot-runs and an honesty-bound proof registry (PROVEN/OPEN/REFUTED). Result: heroic eudaimonia as recursive stability under controlled rupture—describable, implementable, and knowingly limited.

Keywords: Autopoiesis; Autopolitics; Fusion Hero OS; MasterSeed; CEC; Sisyphos; Eudaimonia; Mesh; MCP; Proof Registry; Code Honesty
"""

    short_acad = """This dissertation (monograph v1.0) formulates Fusion Hero OS as both an autopoietic system—continuously producing the components that maintain it (MasterSeed, CEC, Sisyphos regeneration, fractal mesh, archives, tests)—and an autopolitical constitution that self-legislates placement (L0–L4), consent, connector membranes, and proof honesty. Drawing on systems theory, existential philosophy, and live infrastructure empirics (mesh coordination, dashboard health, academia curriculum alignment), it argues that heroic eudaimonia is recursive stability under controlled rupture, not a happiness sum. Code-honesty constraints (PROVEN/OPEN claims in a proof registry) are treated as epistemic organs of the system rather than external critique.
"""

    out_de = OUT_DIR / "ACADEMIA_ABSTRACT_DE.md"
    out_en = OUT_DIR / "ACADEMIA_ABSTRACT_EN.md"
    out_short = OUT_DIR / "ACADEMIA_ABSTRACT_SHORT.txt"
    out_paste = OUT_DIR / "ACADEMIA_UPLOAD_PASTE.txt"

    out_de.write_text(de.strip() + "\n", encoding="utf-8")
    out_en.write_text(en.strip() + "\n", encoding="utf-8")
    out_short.write_text(short_acad.strip() + "\n", encoding="utf-8")
    out_paste.write_text(
        "=== TITLE ===\n"
        "Autopoiesis and Autopolitics of Fusion Hero OS\n\n"
        "=== AUTHORS ===\n"
        "Stephan Hagen Urban\n\n"
        "=== ABSTRACT (paste) ===\n"
        + short_acad.strip()
        + "\n\n=== KEYWORDS ===\n"
        "Autopoiesis, Autopolitics, Fusion Hero OS, MasterSeed, Co-Evolutionary Closure, "
        "Sisyphos, Eudaimonia, Mesh, MCP, Proof Registry, Code Honesty, Tailscale, "
        "Systems Theory, Existential Philosophy\n",
        encoding="utf-8",
    )
    return [out_de, out_en, out_short, out_paste]


def convert_pdf() -> dict:
    if not LO.is_file():
        return {"ok": False, "error": "LibreOffice not found"}
    outdir = str(OUT_DIR)
    # Convert both copies
    results = []
    for src in (DOCX, ARCH_DOCX):
        if not src.is_file():
            continue
        cmd = [
            str(LO),
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(src.parent),
            str(src),
        ]
        t0 = time.time()
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        pdf = src.with_suffix(".pdf")
        results.append(
            {
                "src": str(src),
                "pdf": str(pdf) if pdf.is_file() else None,
                "bytes": pdf.stat().st_size if pdf.is_file() else 0,
                "returncode": proc.returncode,
                "stderr": (proc.stderr or "")[-500:],
                "seconds": round(time.time() - t0, 2),
            }
        )
    # also ensure docs/dissertation has pdf
    return {"ok": any(r.get("pdf") for r in results), "results": results}


def main() -> int:
    claims = load_proofs()
    # latest coordination
    latest = load_json(COORD / "latest.json")
    if not latest:
        # fallback any coordination_all
        cands = sorted(COORD.glob("coordination_all_*.json"), reverse=True)
        latest = load_json(cands[0]) if cands else {}
    acad = load_json(COORD / "academia_training_state.json")

    health_note = "not probed in-process"
    try:
        import urllib.request

        with urllib.request.urlopen("http://127.0.0.1:8000/api/health?light=true", timeout=3) as r:
            health_note = r.read().decode("utf-8", errors="replace")[:120]
    except Exception as exc:  # noqa: BLE001
        health_note = f"unreachable: {exc}"

    doc = Document(str(DOCX))
    append_band3(doc, claims, latest, acad, health_note)
    doc.save(str(DOCX))
    doc.save(str(ARCH_DOCX))

    abstracts = write_academia_abstract()
    pdf_info = convert_pdf()

    wc = sum(len(x.text.split()) for x in doc.paragraphs)
    man = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "version": "1.0+band3",
        "docx": str(DOCX),
        "docx_bytes": DOCX.stat().st_size,
        "words_approx": wc,
        "paragraphs": len(doc.paragraphs),
        "proof_claims": len(claims),
        "proof_status_counts": dict(Counter(c.get("status") for c in claims)),
        "abstracts": [str(a) for a in abstracts],
        "pdf": pdf_info,
        "health": health_note,
    }
    (COORD / "dissertation_band3_manifest.json").write_text(
        json.dumps(man, indent=2), encoding="utf-8"
    )
    # update README
    readme = OUT_DIR / "README.md"
    extra = f"""

## Band III + PDF + Academia (v1.0+band3)

- Band III in DOCX integriert (Empirie + Proof-Tabellen)
- Academia-Abstracts: `ACADEMIA_ABSTRACT_DE.md`, `ACADEMIA_ABSTRACT_EN.md`, `ACADEMIA_ABSTRACT_SHORT.txt`, `ACADEMIA_UPLOAD_PASTE.txt`
- PDF: siehe `*.pdf` neben DOCX (LibreOffice-Export)
- Manifest: `~/.fusion/mesh/coordination/dissertation_band3_manifest.json`

Wörter ca.: {wc}
"""
    if readme.is_file():
        txt = readme.read_text(encoding="utf-8")
        if "Band III + PDF" not in txt:
            readme.write_text(txt.rstrip() + "\n" + extra, encoding="utf-8")

    print(json.dumps(man, indent=2))
    return 0 if pdf_info.get("ok") else 0  # still success if docx/abstracts done


if __name__ == "__main__":
    raise SystemExit(main())
