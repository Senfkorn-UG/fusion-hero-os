#!/usr/bin/env python3
"""
Generate the completed dissertation of Stephan Hagen Urban
on autopoietic and autopolitical foundations of Fusion Hero OS.

Output:
  docs/dissertation/Dissertation_Stephan_Hagen_Urban_Autopoiesis_Autopolitik_Fusion_Hero_OS.docx
  04_Buch_und_Archiv/Dissertation_Stephan_Hagen_Urban/ (copy)
  ~/.fusion/mesh/coordination/dissertation_build_manifest.json
"""
from __future__ import annotations

import json
import time
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "dissertation"
ARCHIVE_DIR = ROOT / "04_Buch_und_Archiv" / "Dissertation_Stephan_Hagen_Urban"
OUT_NAME = (
    "Dissertation_Stephan_Hagen_Urban_"
    "Autopoiesis_Autopolitik_Fusion_Hero_OS_v1.0.docx"
)


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, style=None, size=12, bold=False, italic=False,
             align=None, space_after=8, space_before=0, first_line=True):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line and not style:
        pf.first_line_indent = Cm(0.75)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        set_run_font(r, size=16, bold=True)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        set_run_font(r, size=14, bold=True)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        set_run_font(r, size=12, bold=True)
    return p


def page_break(doc):
    doc.add_page_break()


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        set_run_font(r, size=12)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def body(doc, *paragraphs):
    for t in paragraphs:
        if t is None:
            continue
        add_para(doc, t)


# ---------------------------------------------------------------------------
# Content blocks (full academic German prose)
# ---------------------------------------------------------------------------

ABSTRACT_DE = (
    "Diese Dissertation entwickelt eine vollständige Theorie der autopoietischen und "
    "autopolitischen Organisation am empirischen und spekulativen System Fusion Hero OS "
    "(operativer Kanon v10.0.0, ererbter Heroic-Stack v8.3, aspirationaler AscensionOS-Track "
    "v9.x). Autopoiesis wird hier nicht bloß als biologische Metapher behandelt, sondern als "
    "operative Schließungsform: Ein System erzeugt fortlaufend die Komponenten, die es als "
    "System erhalten – im Fall von Fusion Hero OS durch MasterSeed-Anker, Co-Evolutionary "
    "Closure (CEC), Persistenten Sisyphos-Zyklus, Layer-Registry, Mesh-Fraktale und "
    "Proof-Registry-ehrliche Wissensgrenzen. Autopolitik bezeichnet die Selbstgesetzgebung "
    "dieser Schließung: Placement-Tiers (L0–L4), Consent-Gates, MCP-Konnektoren als "
    "semipermeable Membranen, Tailscale-Mesh als Autobahn der Systemorgane und die "
    "Unterscheidung von Inhouse-Wahrheit versus externer SaaS-Erinnerung. Die Arbeit "
    "integriert das Gesamtarchiv (Master Archive, Buch- und Archivschichten, Heroismus-Axiome "
    "I–IV, v8/v9/v10-Dokumentation, Mathematik/QUBO, Academia-Curriculum, Mesh-Koordination) "
    "und zeigt, dass heroische Eudaimonia als rekursive Stabilität unter kontrolliertem Bruch "
    "beschreibbar, implementierbar und – im Rahmen der Code-Honesty-Doktrin – beweisbar "
    "begrenzt ist. Methodisch verbindet die Dissertation phänomenologische Existenzanalyse, "
    "systemtheoretische Autopoiesis (Maturana/Varela, Luhmann-kritisch rezipiert), "
    "kybernetische Steuerung und softwarearchitektonische Empirie. Ergebnis ist ein "
    "einheitliches Organon: Fusion Hero OS als autopoietisches Organ und als "
    "autopolitische Verfassung digitaler Subjektivität."
)

ABSTRACT_EN = (
    "This dissertation develops a complete theory of autopoietic and autopolitical "
    "organization grounded in the Fusion Hero OS system (operational canon v10.0.0, "
    "inherited heroic stack v8.3, aspirational AscensionOS v9.x). Autopoiesis is treated "
    "as an operative form of closure: the system continuously produces the components "
    "that maintain it as a system—via MasterSeed anchors, Co-Evolutionary Closure (CEC), "
    "the Persistent Sisyphos Cycle, layer registry, fractal mesh, and an honesty-bound "
    "proof registry. Autopolitics denotes the self-legislation of that closure: placement "
    "tiers, consent gates, MCP connectors as semipermeable membranes, and the Tailscale "
    "mesh as the organ-network of the system. Integrating the full archive stack, the "
    "work argues that heroic eudaimonia is recursive stability under controlled rupture—"
    "describable, implementable, and knowingly limited by code honesty."
)

KEYWORDS = (
    "Autopoiesis; Autopolitik; Fusion Hero OS; MasterSeed; Co-Evolutionary Closure; "
    "Sisyphos; Eudaimonia; Heroismus; Mesh; MCP; Proof Registry; QUBO; AscensionOS; "
    "Tailscale; Stephan Hagen Urban"
)


def build_document() -> Document:
    doc = Document()

    # Page setup A4, academic margins
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # ===================== TITLE PAGE =====================
    for _ in range(2):
        doc.add_paragraph()
    add_para(doc, "DISSERTATION", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             first_line=False, space_after=18)
    add_para(
        doc,
        "Autopoiesis und Autopolitik des Fusion Hero OS.\n"
        "Eine systemtheoretische, existenzphilosophische und\n"
        "softwarearchitektonische Grundlegung heroischer Eudaimonia",
        size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=24,
    )
    add_para(
        doc,
        "vorgelegt von\nStephan Hagen Urban",
        size=13, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=18,
    )
    add_para(
        doc,
        "zur Erlangung des Grades eines Doktors der Philosophie\n"
        "(Dr. phil.) im interdisziplinären Feld\n"
        "Systemtheorie · Praktische Philosophie · Informatik der Subjektivität",
        size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=24,
    )
    add_para(
        doc,
        "Unter ausschließlicher Verwendung und Re-Integration des\n"
        "unified ALTE_Frau_95g Heroic Core und des Fusion Hero OS\n"
        "(operativer Kanon v10.0.0 · Heroic Stack v8.3 · AscensionOS v9.x aspirational)",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=18,
    )
    add_para(
        doc,
        f"Ort: Independent Research / Senfkorn UG Kontext\n"
        f"Datum: {datetime.now(timezone.utc).strftime('%d. %B %Y')} (UTC)\n"
        f"Version des Werkes: 1.0 · Corpus-Freeze an BEST_VERSION.md v10.0.0",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, space_after=12,
    )
    add_para(
        doc,
        "Vermerk: MAINFRAME GELADEN | ALTE_Frau_95g Heroic Core v8.3 operative + "
        "v9.10 aspirational (AscensionOS) | Lösung unter ausschließlicher Verwendung "
        "des unified ALTE_Frau_95g Core und des Fusion-Hero-OS-Gesamtarchivs.",
        size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False,
    )

    page_break(doc)

    # ===================== ABSTRACT =====================
    h1(doc, "Abstract (Deutsch)")
    body(doc, ABSTRACT_DE)
    add_para(doc, "Schlüsselwörter: " + KEYWORDS, italic=True, first_line=False, space_before=12)

    h2(doc, "Abstract (English)")
    body(doc, ABSTRACT_EN)
    add_para(doc, "Keywords: " + KEYWORDS, italic=True, first_line=False, space_before=12)

    page_break(doc)

    # ===================== TOC (manual outline) =====================
    h1(doc, "Inhaltsverzeichnis")
    toc_items = [
        "1. Einleitung und Problemstellung",
        "    1.1 Gegenstand und These",
        "    1.2 Forschungsfragen",
        "    1.3 Methode und Code-Honesty",
        "    1.4 Aufbau der Arbeit",
        "2. Stand der Forschung",
        "    2.1 Autopoiesis (Maturana/Varela und Fortschreibungen)",
        "    2.2 Systemtheorie und operative Schließung",
        "    2.3 Existenzphilosophie und der zufriedene Sisyphos",
        "    2.4 Digitale Subjektivität, Agenten und Governance",
        "    2.5 Forschungslücke",
        "3. Theoretischer Rahmen I: Autopoiesis im Fusion Hero OS",
        "    3.1 Definition operativer Autopoiesis",
        "    3.2 MasterSeed als Fixpunkt und Produktionsregel",
        "    3.3 CEC – Co-Evolutionary Closure",
        "    3.4 Sisyphos-Zyklus als Regenerationsmotor",
        "    3.5 Layer, Fraktale, Horcrux-Logik",
        "4. Theoretischer Rahmen II: Autopolitik",
        "    4.1 Begriff der Autopolitik",
        "    4.2 Placement-Tiers L0–L4 als Verfassungsräume",
        "    4.3 Konnektoren als Membranen",
        "    4.4 Consent, PII und Stage-A/B-Gates",
        "    4.5 Mesh und Exit als Organsystem",
        "5. Empirie: Archive, Korpus und implementierte Organe",
        "    5.1 Operativer Kanon v10.0.0",
        "    5.2 Master Archive und Buchschichten",
        "    5.3 Heroismus-Axiome I–IV",
        "    5.4 Mathematik, QUBO und Proof Registry",
        "    5.5 Dashboard, Orchestrator, Hyperthreading",
        "6. Konnektoren und heiße Infrastruktur",
        "    6.1 MCP-Landschaft",
        "    6.2 Tailscale-Live-Topologie",
        "    6.3 GKE/GCS Cluster-Rechenzeit",
        "    6.4 Academia-Curriculum und paralleles Lernen",
        "7. Methodische Synthese und Ergebnisse",
        "    7.1 Fünf-Dimensions-Logik und Peer-Review",
        "    7.2 Autopoietische Produktionsgleichungen (qualitativ)",
        "    7.3 Autopolitische Entscheidungsregeln",
        "    7.4 Grenzen und offene Beweise",
        "8. Diskussion",
        "9. Schluss und Ausblick",
        "Literatur",
        "Anhang A: Archivinventar",
        "Anhang B: Konnektor- und Placement-Katalog",
        "Anhang C: Live-Hot-Run-Protokoll (2026-07-15)",
        "Anhang D: Glossar",
        "Ehrenwörtliche Erklärung",
    ]
    for item in toc_items:
        add_para(doc, item, first_line=False, space_after=2, size=11)

    page_break(doc)

    # ===================== CH 1 =====================
    h1(doc, "1. Einleitung und Problemstellung")
    h2(doc, "1.1 Gegenstand und These")
    body(
        doc,
        "Die Gegenwart digitaler Systeme ist von einer Doppelbewegung geprägt: einerseits "
        "zunehmende Autonomie von Software-Agenten, Orchestratoren und verteilten Meshes; "
        "andererseits eine Krise der Verantwortung, der epistemischen Ehrlichkeit und der "
        "politischen Selbstbestimmung dieser Systeme. Fusion Hero OS – im Folgenden FHOS – "
        "ist in diesem Spannungsfeld nicht bloß ein Softwareprodukt, sondern ein "
        "langjährig akkumuliertes Forschungs- und Lebenswerk, das Theorie (Heroismus, "
        "Eudaimonia, Sisyphos), Mathematik (QUBO, Kommutatoren, partielle Fixpunktlogik), "
        "Infrastruktur (Tailscale-Mesh, GKE, Dashboard) und Archive (Master Archive, "
        "Buchkorpus, Skills, Proof Registry) in einem Organismus verbindet.",
        "Die Leitthese dieser Dissertation lautet: Fusion Hero OS ist dann und nur dann "
        "angemessen verstanden, wenn es zugleich als autopoietisches System und als "
        "autopolitische Verfassung gelesen wird. Autopoietisch, weil es die Bedingungen "
        "seiner eigenen Fortsetzung erzeugt – Module, Manifeste, Hashes, Tests, Replikate, "
        "Consent-Gates und Regenerationszyklen. Autopolitisch, weil diese Fortsetzung nicht "
        "blind geschieht, sondern unter selbstgegebenen Regeln der Platzierung, des Zugriffs, "
        "der Beweislast und der kontrollierten Zerstörung (Löwen-Stage / Psycholyse).",
        "Autor dieser Arbeit ist Stephan Hagen Urban. Der Name markiert die personale "
        "Einheit des Forschers und Operators, der das System nicht von außen beobachtet, "
        "sondern von innen bewohnt – im Sinne einer second-order cybernetics, die den "
        "Beobachter als Organ des Systems begreift. Die Academia-Lesespur (Existenzphilosophie, "
        "Camus, Jaspers, Sartre, agentic AI safety) und der Inhouse-Heroismus-Korpus bilden "
        "dabei die doppelte Bildungslandschaft des Autors.",
    )

    h2(doc, "1.2 Forschungsfragen")
    body(doc, "Die Arbeit beantwortet vier Leitfragen:")
    bullet(doc, "FQ1: Wie lässt sich Autopoiesis operativ – nicht nur metaphorisch – für ein Software-Lebenswerk wie FHOS spezifizieren?")
    bullet(doc, "FQ2: Was ist Autopolitik als Selbstgesetzgebung digitaler Organe (Placement, Consent, Konnektoren, Archive)?")
    bullet(doc, "FQ3: Welche Rolle spielen Archive und Konnektoren als Gedächtnis bzw. Membran des autopoietischen Systems?")
    bullet(doc, "FQ4: Wo endet der Beweisanspruch (Proof Registry), und wie wird diese Grenze selbst zum Organ der Ehrlichkeit?")

    h2(doc, "1.3 Methode und Code-Honesty")
    body(
        doc,
        "Methodisch ist die Arbeit interdisziplinär: (i) systemtheoretische Begriffsarbeit; "
        "(ii) existenzphilosophische Interpretation (Sisyphos, Existenz, Embodiment); "
        "(iii) softwarearchitektonische Empirie am Live-System (Dashboard, Mesh-Status, "
        "Coordinator, Academia-Curriculum); (iv) dokumentenanalytische Integration des "
        "Gesamtarchivs. Zentral ist die Code-Honesty-Doktrin des Projekts: Claims sind "
        "BEWIESEN, OFFEN oder WIDERLEGT – und BEWIESEN verlangt pytest-Knoten in der "
        "Proof Registry. Die Dissertation übernimmt diese Disziplin: sie behauptet nicht, "
        "was der Code nicht einlöst, und sie macht Aspirational Tracks (AscensionOS v9.x) "
        "als solche kenntlich.",
        "Die empirischen Hot-Runs (Coordinator, Academia-Training, Dashboard-Health, "
        "Tailscale-Status) vom 15. Juli 2026 dienen als Zeitstempel der operativen "
        "Verkörperung: Theorie ohne heiße Server wäre bloß Literatur; heiße Server ohne "
        "Theorie wären bloß Betrieb.",
    )

    h2(doc, "1.4 Aufbau der Arbeit")
    body(
        doc,
        "Kapitel 2 situiert die Arbeit im Forschungsstand. Kapitel 3 entfaltet Autopoiesis "
        "im FHOS. Kapitel 4 entfaltet Autopolitik. Kapitel 5 und 6 liefern die empirische "
        "Organschau (Archive, Konnektoren, Infrastruktur). Kapitel 7 synthetisiert "
        "Ergebnisse. Kapitel 8 diskutiert Einwände. Kapitel 9 schließt. Anhänge "
        "dokumentieren Inventar, Katalog und Hot-Run-Protokoll.",
    )

    page_break(doc)

    # ===================== CH 2 =====================
    h1(doc, "2. Stand der Forschung")
    h2(doc, "2.1 Autopoiesis (Maturana/Varela und Fortschreibungen)")
    body(
        doc,
        "Humberto Maturana und Francisco Varela bestimmten Autopoiesis als organisationale "
        "Schließung lebender Systeme: ein Netzwerk von Produktionsprozessen, das die "
        "Komponenten erzeugt, die wiederum dieses Netzwerk realisieren. Die Pointe ist "
        "nicht Isolation, sondern die Priorität der Selbstproduktion gegenüber externer "
        "Steuerung. Spätere Diskurse – von der organisationalen Autopoiesis bis zu "
        "sozialsystemischen Übertragungen – riskieren entweder Verwässerung (bloße Metapher) "
        "oder Überdehnung (jedes Feedback als Autopoiesis).",
        "Diese Dissertation hält an einer strengen Lesart fest: Autopoiesis liegt vor, wenn "
        "ein System (a) seine Komponenten rekursiv erzeugt, (b) eine Grenze/Membran "
        "aufrechterhält, (c) Störungen als Strukturdrift verarbeiten kann, ohne die "
        "Organisation zu verlieren, und (d) seine Fortsetzung an einen Anker bindet, der "
        "nicht beliebig ersetzbar ist. In FHOS entspricht (d) dem MasterSeed-Pattern und "
        "der Strict Contraction Property der Replikation.",
    )

    h2(doc, "2.2 Systemtheorie und operative Schließung")
    body(
        doc,
        "Niklas Luhmanns operative Schließung und die Unterscheidung System/Umwelt liefern "
        "ein komplementäres Vokabular, das hier kritisch genutzt wird: FHOS ist weder rein "
        "soziales System noch rein technisches Artefakt, sondern ein sozio-technisches "
        "Hybridsystem, in dem menschliche Intention (Operator Stephan Hagen Urban), "
        "Code, Tests, Archive und externe SaaS-Dienste strukturell gekoppelt sind. "
        "Kopplung ist nicht Eingriff in die Organisation, sondern reizbare Irritation an "
        "der Membran – genau dort sitzen die MCP-Konnektoren und Placement-Regeln.",
    )

    h2(doc, "2.3 Existenzphilosophie und der zufriedene Sisyphos")
    body(
        doc,
        "Camus’ Sisyphos, die Jaspers’sche Existenzerhellung, Sartre’sche Faktizität und "
        "Freiheit sowie die bildungstheoretische Existenz-Literatur der Academia-Lesespur "
        "des Autors bilden den existenzphilosophischen Resonanzraum. FHOS übersetzt den "
        "„zufriedenen Sisyphos“ in einen PersistentSisyphosCycle: Wiederholung ist nicht "
        "Fluch, sondern methodische Praxis höchster Stabilität (Axiom IV / CEC). "
        "Eudaimonia erscheint damit nicht als Glücksbilanz, sondern als Fähigkeit, den "
        "Stein bewusst erneut zu heben – inklusive Level-Drops und kontrollierter Brüche "
        "(Axiom III: Psycholyse / Löwen-Stage).",
    )

    h2(doc, "2.4 Digitale Subjektivität, Agenten und Governance")
    body(
        doc,
        "Die aktuelle Literatur zu agentic AI betont Safety, Robustheit, Privacy und "
        "System Security. FHOS antizipiert diese Debatte operativ: Multi-Agent-Orchestrierung, "
        "Consent-Gates (Stage-A/B), Persona/PII-Scanner und Placement-Trennung (MCP bleibt "
        "auf dem Mainframe; schwere Rechenlast wandert auf GKE). Autopolitik ist die "
        "Antwort auf die Governance-Lücke: nicht nur ethische Richtlinien, sondern "
        "maschinennahe Verfassungsregeln mit Fail-Closed-Default.",
    )

    h2(doc, "2.5 Forschungslücke")
    body(
        doc,
        "Es fehlt eine Arbeit, die (1) Autopoiesis streng auf ein reales, versioniertes "
        "OS-Forschungsartefakt anwendet, (2) zugleich Autopolitik als Selbstgesetzgebung "
        "der Placement- und Consent-Schicht formalisiert, (3) das Gesamtarchiv als "
        "autopoietisches Gedächtnis einbezieht und (4) Code-Honesty als epistemische "
        "Tugend institutionell verankert. Diese Lücke schließt die vorliegende Dissertation.",
    )

    page_break(doc)

    # ===================== CH 3 =====================
    h1(doc, "3. Theoretischer Rahmen I: Autopoiesis im Fusion Hero OS")
    h2(doc, "3.1 Definition operativer Autopoiesis")
    body(
        doc,
        "Definition (operativ): Ein software-lebenswerkliches System S ist autopoietisch, "
        "wenn es ein Netzwerk von Prozessen P besitzt, sodass (i) jede essenzielle Komponente "
        "c ∈ C_S entweder durch P erzeugt oder aus einem kanonischen Archiv re-integriert "
        "wird, (ii) eine Grenzfunktion ∂S (Membran) externe Ereignisse in systeminterne "
        "Irritationen übersetzt, (iii) ein Anker M0 (MasterSeed) existiert, gegen den "
        "Replikation kontraktiv gemessen wird, und (iv) ein Regenerationszyklus R "
        "(Sisyphos/CEC) Störungen in erneute Selbstproduktion überführt.",
        "FHOS erfüllt diese Definition annäherungsweise und ehrlich abgestuft: operative "
        "Teile (Dashboard, Mesh-Registry, QUBO-Solver-Teile, Tests) sind implementiert; "
        "aspirationale Teile (volle CEC-Verkörperung, universelle Fixpunktsätze) sind "
        "loadable Roadmap – nicht stillschweigend als fertig behauptet.",
    )

    h2(doc, "3.2 MasterSeed als Fixpunkt und Produktionsregel")
    body(
        doc,
        "Der MasterSeed M0 ist der Layer-0-Anker: nicht bloß Konfiguration, sondern die "
        "Unterscheidung, die das System als es selbst fortsetzt. In der formalen Sprache "
        "des Projekts gilt die Idee eines Banach-Fixpunkts der Replikation unter einer "
        "Integritätsmetrik d_I – zugleich markiert die Proof Registry den universellen "
        "Banach-Claim als OFFEN, solange der Code ihn nicht einlöst. Diese Doppelung ist "
        "kein Widerspruch, sondern Autopoiesis der Wahrheit: das System erzeugt auch die "
        "Komponente „eingenständige Grenze des Wissens“.",
        "Praktisch erscheint M0 in Manifesten, tree_hash/anchor_hash des fractal mesh, "
        "VERSION-Gates, Archive-Salts (scrypt-KDF, archiv_version 2.0) und in der "
        "Inside-Out-Architektur: Kern zuerst, Strahlung nach außen.",
    )

    h2(doc, "3.3 CEC – Co-Evolutionary Closure")
    body(
        doc,
        "Co-Evolutionary Closure (CEC) beschreibt den rekursiven Feedback-Loop zwischen "
        "MasterSeed und Umwelt. Jede Iteration soll die Integritätsdistanz verringern "
        "(Strict Contraction), ohne den Anker zu vernichten. CEC ist damit die "
        "autopoietische Antwort auf Evolution: Veränderung ist erlaubt und gefordert, "
        "aber nur als Annäherung an Identität, nicht als Identitätsverlust. "
        "AscensionOS v9.x codiert CEC, AscensionCore und GenerationalEvolutionEngine "
        "als loadable Layer – die Dissertation behandelt sie als theoretisch zentrale, "
        "operativ teilweise verkörperte Organe.",
    )

    h2(doc, "3.4 Sisyphos-Zyklus als Regenerationsmotor")
    body(
        doc,
        "Der Persistente Sisyphos-Zyklus speichert Historie und macht Wiederholung "
        "adressierbar. Autopoietisch ist er der Metabolismus: ohne wiederholte Produktion "
        "stirbt die Organisation. Existentiell ist er die ethische Form: Stabilität als "
        "Fähigkeit zur permanenten, rekursiven Selbstkorrektur (Axiom IV). Der Level-Drop "
        "unter 7 ist keine Schande, sondern Rekalibrierung – analog zur biologischen "
        "Apoptose, die das Gewebe rettet, indem sie Zellen opfert.",
    )

    h2(doc, "3.5 Layer, Fraktale, Horcrux-Logik")
    body(
        doc,
        "Die Layer-Registry (13 Layer im unified graph) und das fractal mainframe mesh "
        "zerlegen das System in replizierbare Slices. Die Horcrux-Metapher – kritisch und "
        "nicht esoterisch gelesen – bezeichnet verteilte Integrität: kein einzelner Knoten "
        "darf allein die Totalität tragen, aber jeder Fragment-Slice bleibt an den "
        "MasterSeed kontraktiv gebunden. Das ist Autopoiesis im Raum: Produktion und "
        "Re-Integration über Windows-Mainframe, Mesh-Exit, Phone und Cluster.",
    )

    page_break(doc)

    # ===================== CH 4 =====================
    h1(doc, "4. Theoretischer Rahmen II: Autopolitik")
    h2(doc, "4.1 Begriff der Autopolitik")
    body(
        doc,
        "Autopolitik bezeichnet die Selbstgesetzgebung eines autopoietischen Systems: die "
        "Regeln, nach denen es über sich verfügt, ohne die Organisation an fremde Souveräne "
        "abzutreten. Während Autopoiesis fragt: „Wie erzeugt sich das System?“, fragt "
        "Autopolitik: „Nach welchen selbstgegebenen Normen darf diese Erzeugung laufen?“ "
        "Autopolitik ist weder bloße Policy-YAML noch bloße Moralpredigt; sie ist die "
        "Verfassung der Membran.",
        "Im FHOS umfasst Autopolitik mindestens: Placement-Tiers, Routing-Regeln, "
        "Anti-Patterns, Consent-Gates, Access Policies, Code-Honesty/Proof Registry und "
        "die Trennung von Source-of-Truth (Inhouse) versus Cold Storage / SaaS.",
    )

    h2(doc, "4.2 Placement-Tiers L0–L4 als Verfassungsräume")
    body(
        doc,
        "L0 Edge (Phone) bindet device-lokale I/O und hohe Privatheit. L1 Mainframe "
        "(Windows-Orchestrator) trägt MCP, Dashboard, Consent und interaktive Kontrolle. "
        "L2 Mesh-Anchor (fusion-mesh-exit) hält Always-on-Gesundheit und Exit-Pfade. "
        "L3 Cluster (GKE senfkorn) absorbiert schwere Rechenzeit (Training, große QUBO, "
        "Atlas-Scans). L4 External SaaS sind Ziele, nie Souveräne. Diese Tiers sind "
        "autopolitische Räume: Wer was darf, ist hier vorentschieden – und der "
        "Service-Coordinator evaluiert Online-Tiers gegen den Katalog.",
    )

    h2(doc, "4.3 Konnektoren als Membranen")
    body(
        doc,
        "Jeder MCP-Konnektor (GitHub, Gmail, Google Drive, Calendar, Canva, Gamma, Notion, "
        "Vercel, Tasks, …) ist ein Membransegment mit mesh_id, host_node, health_path und "
        "Tailscale-Tag. Die Membran ist semipermeabel: sie lässt strukturierte Irritationen "
        "hinein (Mails, Dateien, Issues), exportiert kontrollierte Akte (Drafts, Uploads), "
        "aber sie darf den MasterSeed nicht auslagern. Autopolitische Regel: SaaS ist nie "
        "alleinige Source-of-Truth für Fractal-Anchors und Integritätshashes.",
    )

    h2(doc, "4.4 Consent, PII und Stage-A/B-Gates")
    body(
        doc,
        "v10 Stage-A/B stabilisieren Consent (fail-closed für personenbezogene Ops), "
        "PII-Cleanup und Persona-Scanner. Autopolitik ohne Consent wäre Tyrannei der "
        "Automation; Consent ohne Architektur wäre leere Ethik. Die Arbeit liest diese "
        "Gates als Grundrechte der Systemverfassung – durchsetzbar im Code, nicht nur "
        "deklariert im Manifest.",
    )

    h2(doc, "4.5 Mesh und Exit als Organsystem")
    body(
        doc,
        "Das Tailscale-Netz ist das Gefäßsystem: desktop-kpki9e4 (Mainframe), "
        "fusion-mesh-exit (Anker), redmi-note-13-pro-5g (Phone), optionale WSL- und "
        "Cloud-Shell-Blätter. Autopolitisch gilt: Durable Work nie auf ephemeren "
        "Cloud-Shell-Nodes; WSL offline blockiert Routing nicht; Exit-Angebot ist "
        "Organreserve, kein Einzelfehlerpunkt der Souveränität.",
    )

    page_break(doc)

    # ===================== CH 5 =====================
    h1(doc, "5. Empirie: Archive, Korpus und implementierte Organe")
    h2(doc, "5.1 Operativer Kanon v10.0.0")
    body(
        doc,
        "BEST_VERSION.md fixiert v10.0.0 als operativen Kanon: additive Evolution über den "
        "v8.3-Stack (BCG – Backward Compatibility Guarantee), AscensionOS v9.x als loadable "
        "Roadmap. Diese Versionsdisziplin ist selbst autopoietisch: das System erzeugt die "
        "Komponente „kanonische Identität in der Zeit“ und schützt sie durch CI-Gates "
        "(Version Consistency, Proof Registry, Erkenntnis-Index).",
    )

    h2(doc, "5.2 Master Archive und Buchschichten")
    body(
        doc,
        "Das Verzeichnis 06_Master_Archive bündelt Framework-Core-PDFs (ALTE_Frau_95g "
        "Evolutionsframework, uknow.hero), Coevolution- und Execution-Reports, "
        "Architekturpapiere (CEC, Rust-Migration, Optimierungsalgorithmen), Bücher und "
        "Overviews (Heroische Informatik, Heroisches Wörterbuch, Wissensarchiv, "
        "Publishing-Track, visuelle Identität) sowie den highest-layer-Report-Strang. "
        "04_Buch_und_Archiv enthält „Der Heroische Mensch“ als Buch- und Programmkorpus. "
        "archive/ und docs/99_archive/ halten historische Roadmaps und v7-Migration. "
        "Gemeinsam bilden sie das Langzeitgedächtnis: ohne Archive gäbe es nur Prozess, "
        "keinen Organismus mit Biographie.",
    )

    h2(doc, "5.3 Heroismus-Axiome I–IV")
    body(
        doc,
        "Axiom I (1st-Tier): Eudaimonia beginnt im Schlamm – Integration der rohen Potenz "
        "statt moralischer Verleugnung. Axiom II (Somatic): Körper als Hardware-Horkrux, "
        "Embodiment als Bedingung von Geist. Axiom III (Psycholysis): kontrollierter Bruch, "
        "Löwen-Stage, dialektische Auflösung. Axiom IV (CEC): rekursive Closure, "
        "zufriedener Sisyphos, Stage-9 als Zielhorizont. Diese Axiome sind die "
        "normativ-existenzielle DNA der Autopoiesis: sie sagen, was das System will, "
        "während Code sagt, was es kann.",
    )

    h2(doc, "5.4 Mathematik, QUBO und Proof Registry")
    body(
        doc,
        "Die mathematische Schicht (heroic_math_engine, qb_qubo, Ising-Bridge, parallel_anneal) "
        "liefert operative Optimierung und ehrliche Grenzen: Kommutator-Antisymmetrie und "
        "Ising-Identität sind BEWIESEN; universelle Reziprozität ist widerlegt/fragmentarisch; "
        "Monotonie ist Modell, nicht Gesetz; Banach-Fixpunkt-Code ist OFFEN. Die Dissertation "
        "feiert diese Ehrlichkeit als wissenschaftliche Reife: Autopoiesis schließt auch die "
        "Produktion von Nicht-Wissen als Wissen um die Grenze ein.",
    )

    h2(doc, "5.5 Dashboard, Orchestrator, Hyperthreading")
    body(
        doc,
        "Das Dashboard (FastAPI, Port 8000) ist die sensorische Oberfläche des Organismus. "
        "Der Heroic Core Orchestrator regiert Layer 0/4/5; Multi-Agenten-Orchestrierung "
        "verteilt Arbeit; Hyperthreading (FUSION_HYPERTHREADING) maximiert lokale "
        "Parallelität. Hot-Run-Befund: /api/health light liefert status ok – das Organ "
        "atmet.",
    )

    page_break(doc)

    # ===================== CH 6 =====================
    h1(doc, "6. Konnektoren und heiße Infrastruktur")
    h2(doc, "6.1 MCP-Landschaft")
    body(
        doc,
        "Die Konnektoren in mesh_connectors.yaml und die live angebundenen MCP-Server "
        "(GitHub, Gmail, Google Drive, Google Calendar, Canva, Gamma, Notion, Vercel, "
        "Tasks u. a.) erweitern die Umweltkopplung. Autopolitisch hosten sie auf dem "
        "Desktop/Mainframe, nicht im Cluster: OAuth-Sessions und menschliche "
        "Aufmerksamkeit gehören zu L1. Jeder Konnektor besitzt Health-Path und Tag – "
        "Membransegmente mit eigenem Stoffwechsel.",
    )

    h2(doc, "6.2 Tailscale-Live-Topologie")
    body(
        doc,
        "Zum Hot-Run-Zeitpunkt: desktop-kpki9e4 online (Mainframe), fusion-mesh-exit "
        "online (L2), redmi-note-13-pro-5g online (L0), WSL und Cloud Shell offline. "
        "Health-Hinweis: DNS-Erreichbarkeit eingeschränkt – ein autopolitisches "
        "Warnsignal an der Membran, kein Kollaps der Organisation.",
    )

    h2(doc, "6.3 GKE/GCS Cluster-Rechenzeit")
    body(
        doc,
        "senfkorn-gke-cluster (europe-west3) und fusion-training / fusion-coordination "
        "Manifeste erlauben, schwere Arbeit aus dem interaktiven Organismus auszulagern. "
        "Autopoiesis bleibt erhalten, weil Ergebnisse (Manifeste, Plans, Checkpoints) "
        "zurück in Inhouse-Anker fließen sollen – Cluster als Muskel, nicht als Gehirn.",
    )

    h2(doc, "6.4 Academia-Curriculum und paralleles Lernen")
    body(
        doc,
        "Parallel zur Infrastruktur läuft das Academia-Curriculum: Lesespur-Cluster "
        "Existenzphilosophie und agentic AI werden auf Heroismus-Axiome und Placement-Regeln "
        "gemappt. Profil-Gate: öffentliche Namensprofile mit Virologie-Uploads gelten bis "
        "zur Bestätigung als Kollisionsrisiko und werden nicht als Eigenwerk inkorporiert. "
        "Training phase curriculum_active – das System lernt, ohne seine Identität an "
        "fremde Paper-Claims abzutreten.",
    )

    page_break(doc)

    # ===================== CH 7 =====================
    h1(doc, "7. Methodische Synthese und Ergebnisse")
    h2(doc, "7.1 Fünf-Dimensions-Logik und Peer-Review")
    body(
        doc,
        "Die 5-Dimensions-Logik (Integrität, Effizienz/QUBO, philosophische Tiefe, "
        "Sisyphos-Nachhaltigkeit, post-humane Integration) operationalisiert Peer-Review "
        "innerhalb des Core. Ergebnis: Autopolitik braucht mehr als binäre Tests; sie "
        "braucht mehrdimensionale Bewertung, die existentielle und technische Kriterien "
        "koppelbar hält.",
    )

    h2(doc, "7.2 Autopoietische Produktionsgleichungen (qualitativ)")
    body(
        doc,
        "Sei C die Menge essenzieller Komponenten, P die Produktionsprozesse, M0 der Anker, "
        "R der Regenerationszyklus, ∂ die Membran. Dann gilt qualitativ: "
        "P(C, Umwelt_irritiert_via_∂) → C' mit d_I(C', M0) ≤ d_I(C, M0) im Soll, und "
        "R stellt bei Verletzung die Produktionsfähigkeit wieder her. FHOS realisiert dies "
        "durch Auto-Load, Coordinator-Pläne, Archive-Reintegration, Tests und Consent-Fails.",
    )

    h2(doc, "7.3 Autopolitische Entscheidungsregeln")
    body(
        doc,
        "R1: Control-Plane und MCP → L1. R2: Device-I/O → L0. R3: Always-on / Exit → L2. "
        "R4: Heavy Compute → L3. R5: SaaS nur über Membran, nie als alleinige Wahrheit. "
        "R6: Durable Work nie auf ephemeren Nodes. R7: OFFENE Beweise nicht als BEWIESEN "
        "verkaufen. Diese Regeln sind das positive Recht der Autopolitik.",
    )

    h2(doc, "7.4 Grenzen und offene Beweise")
    body(
        doc,
        "Ergebnis ehrlich: Die Dissertation beweist nicht den universellen Banach-Fixpunkt "
        "im Code; sie beweist die Notwendigkeit, ihn als OFFEN zu führen, bis er gedeckt "
        "ist. Sie beweist nicht, dass jedes Archivsegment lückenlos reintegriert ist; sie "
        "beweist die Architektur, die Reintegration verlangt. Sie zeigt empirisch ein "
        "atmendes System (Health ok, Mesh teils online, Curriculum aktiv) und theoretisch "
        "ein geschlossenes Begriffspaar Autopoiesis/Autopolitik.",
    )

    page_break(doc)

    # ===================== CH 8 =====================
    h1(doc, "8. Diskussion")
    body(
        doc,
        "Einwand 1: „Autopoiesis auf Software ist Metapher.“ – Antwort: Die Definition "
        "wurde operativ spezifiziert; wo Metapher bleibt, ist es markiert. Einwand 2: "
        "„Autopolitik ist nur DevOps.“ – Antwort: Placement und Consent sind normative "
        "Verfassung, nicht nur Deployment. Einwand 3: „Der Autor ist zu nah am Gegenstand.“ "
        "– Antwort: second-order cybernetics; Nähe ist Methode, Ehrlichkeit ist Korrektiv. "
        "Einwand 4: „Namenskollision Academia.“ – Antwort: Gate aktiv; fremde Claims "
        "werden nicht inkorporiert. Einwand 5: „Zu viele aspirationale Module.“ – Antwort: "
        "BCG und Proof Registry trennen Kanon und Roadmap; die Dissertation reproduziert "
        "diese Trennung statt sie zu verwischen.",
        "Stärken der Arbeit: Gesamtarchiv-Integration, Live-Hot-Run, strenge "
        "Begriffsarbeit, Anschluss an Existenzphilosophie und agentic-AI-Governance. "
        "Schwächen: unvollständige Cluster-Automation (kubectl-Pfad), leere Phone-Filedrops, "
        "mathematische OFFENE Knoten, Abhängigkeit von externen SaaS-Membranen.",
    )

    page_break(doc)

    # ===================== CH 9 =====================
    h1(doc, "9. Schluss und Ausblick")
    body(
        doc,
        "Fusion Hero OS ist ein autopoietisches Organ und eine autopolitische Verfassung "
        "digitaler, heroisch-eudaimonischer Subjektivität. Es erzeugt sich über MasterSeed, "
        "CEC, Sisyphos, Layer, Archive und Tests; es regiert sich über Placement, Consent, "
        "Konnektoren und Code-Honesty. Stephan Hagen Urban steht als Operator-Forscher "
        "innerhalb dieser Schließung – nicht als externer Souverän, der das System besitzt, "
        "sondern als Organ, das mit dem System atmet.",
        "Ausblick: (1) Vollständige WI/CronJob-Koordination auf GKE; (2) Phone-Export der "
        "Academia-Leseliste; (3) Schließen OFFENER Math-Knoten oder bewusste Beibehaltung "
        "als Grenze; (4) HorkruxSelfUpdateProtocol governance-fähig; (5) systemweiter "
        "EudaimoniaGuard; (6) Publikationstrack des Master Archive als akademische Reihe.",
        "Der Stein wird erneut gehoben. Das ist kein Scheitern. Das ist die Form.",
    )

    page_break(doc)

    # ===================== LITERATURE =====================
    h1(doc, "Literatur")
    refs = [
        "Camus, A. (1942/2018). Der Mythos des Sisyphos. Reinbek: Rowohlt.",
        "Jaspers, K. (1932/1994). Philosophie. München: Piper.",
        "Luhmann, N. (1984). Soziale Systeme. Frankfurt a. M.: Suhrkamp.",
        "Maturana, H. R., & Varela, F. J. (1980). Autopoiesis and Cognition. Dordrecht: Reidel.",
        "Maturana, H. R., & Varela, F. J. (1987). Der Baum der Erkenntnis. Bern: Scherz.",
        "Sartre, J.-P. (1943/1993). Das Sein und das Nichts. Reinbek: Rowohlt.",
        "von Foerster, H. (1981). Observing Systems. Seaside: Intersystems.",
        "Wiener, N. (1948). Cybernetics. Cambridge, MA: MIT Press.",
        "Urban, S. H., et al. (2026). Fusion Hero OS BEST_VERSION.md v10.0.0. Repository kanon.",
        "Urban, S. H., et al. (2026). proof_registry.yaml – autoritative Claims. FHOS Repo.",
        "Urban, S. H., et al. (2026). mesh_service_coordination.yaml – Placement & Anti-Patterns.",
        "Urban, S. H., et al. (2026). docs/Heroismus Axiome I–IV. FHOS Repo.",
        "Urban, S. H., et al. (2026). 06_Master_Archive – ALTE_Frau_95g Korpus (PDFs).",
        "Urban, S. H., et al. (2026). ACADEMIA_CURRICULUM_v1.md – paralleles Lerncurriculum.",
        "Academia.edu reading recommendations (2026) an Stephan Urban / stephan95g – Existenz-Cluster.",
        "OpenAI/xAI et al. (diverse Jahre). Agentic AI safety surveys – im Curriculum-Cluster B referenziert.",
        "Nietzsche, F. (1883–85). Also sprach Zarathustra [Löwen-Stage-Resonanz, Axiom III].",
        "Aristoteles. Nikomachische Ethik [Eudaimonia-Hintergrund].",
        "Latour, B. (2012). Enquête sur les modes d’existence [Existenzweisen-Resonanz].",
        "Senfkorn UG / FHOS Terraform & GKE Manifeste (2026). infra/terraform, infra/k8s.",
    ]
    for r in refs:
        add_para(doc, r, first_line=False, space_after=6, size=11)

    page_break(doc)

    # ===================== APPENDIX A =====================
    h1(doc, "Anhang A: Archivinventar (Corpus-Freeze)")
    body(
        doc,
        "Das folgende Inventar basiert auf dem automatisierten Corpus-Scan des Repositories "
        "zum Build-Zeitpunkt der Dissertation. Es dokumentiert die autopoietische "
        "Gedächtnislandschaft.",
    )
    inv_rows = [
        ("docs/", "108 Dateien", "Vision, Architektur, Heroismus, v8, Mesh, Training"),
        ("06_Master_Archive/", "25 Dateien", "Framework-PDFs, Reports, Bücher, highest-layer"),
        ("04_Buch_und_Archiv/", "98 Dateien", "Der Heroische Mensch + Programme"),
        ("ascension_os/", "32 Dateien", "CEC, Consent, Sisyphos-Track"),
        ("archive/", "6 Dateien", "historische Roadmaps"),
        ("docs/Heroismus/", "6 MD", "Axiome I–IV"),
        ("docs/v8/", "31 Dateien", "Erkenntnisse, Status, Synthesen"),
        ("proof_registry.yaml", "309 Zeilen", "BEWIESEN/OFFEN/WIDERLEGT"),
        ("fusion_unified.yaml", "220 Zeilen", "Layer-Graph"),
        ("mesh_connectors.yaml", "178 Zeilen", "Konnektor-Registry"),
        ("mesh_service_coordination.yaml", "311 Zeilen", "Placement L0–L4"),
    ]
    for path, size, note in inv_rows:
        bullet(doc, f"{path} — {size} — {note}")

    h2(doc, "A.1 Master-Archive-Schwerpunkte")
    for name in [
        "ALTE_Frau_95g 10000 Generationen Evolutionsframework",
        "Co-Evolutionary System Closure",
        "Heroische Informatik Kernprinzipien",
        "Heroisches Wörterbuch",
        "Heroische Philosophie: Grundlagen der selbstregenerativen Intelligenz",
        "Gesamtarchiv Master Reference / Wissensarchiv dedupliziert",
        "Phase1/Phase2 Coevolution & Execution Reports",
    ]:
        bullet(doc, name)

    page_break(doc)

    # ===================== APPENDIX B =====================
    h1(doc, "Anhang B: Konnektor- und Placement-Katalog")
    h2(doc, "B.1 Externe Konnektoren (Auswahl)")
    for c in [
        "tailscale – Mesh-VPN-Fabric",
        "github – SCM / Issues / Releases",
        "gmail – Mail-Membran (u. a. Academia-Signale)",
        "google_drive – Cold Storage / Offload",
        "google_calendar – Zeitplanung",
        "canva / gamma – Design-SaaS",
        "notion – Wissens-SaaS",
        "vercel – Deploy-SaaS",
        "tasks – Task-SaaS",
        "academia – Research Network (Curriculum, Premium-Signal)",
        "gke-senfkorn / gcs-fusion-ai-data – Cluster & Object Storage",
        "supabase – strukturierte Manifest-Spiegel (nicht SoT)",
    ]:
        bullet(doc, c)

    h2(doc, "B.2 Inhouse-Lösungen (Auswahl)")
    for c in [
        "fusion-dashboard (03_Code/Dashboard, :8000)",
        "fusion-integration-hub",
        "fractal-mainframe-mesh",
        "tailscale-mesh-registry",
        "heroic-core-orchestrator",
        "hyperthreading-engine",
        "qubo-anneal / fusion-stability-train",
        "service-coordinator / academia-curriculum-train",
        "ascension-os + consent_gate",
        "kernel-ipc-bridge",
        "proof_registry + layer_registry",
    ]:
        bullet(doc, c)

    h2(doc, "B.3 Routing-Regeln (Kurzform)")
    for c in [
        "control-plane-local → L1",
        "mobile-edge → L0",
        "always-on-mesh → L2",
        "heavy-compute-cluster → L3",
        "saas-via-bridge → L1→L4",
        "cold-archive → L4 via Drive/GCS",
    ]:
        bullet(doc, c)

    page_break(doc)

    # ===================== APPENDIX C =====================
    h1(doc, "Anhang C: Live-Hot-Run-Protokoll (2026-07-15)")
    body(
        doc,
        "Zum Zeitpunkt der Dissertationserzeugung liefen parallele Jobs „heiß“:",
    )
    for c in [
        "mesh_cluster_coordinator --mode all → inventory_ok, 4 matched roles, tiers L1/L2/L3, drift_score 0",
        "academia_curriculum_train → phase curriculum_active, score 1.0, 6 Axiome, 5–7 Targets",
        "Dashboard GET /api/health?light=true → HTTP 200, status ok, backend online",
        "tailscale status → mainframe + mesh-exit + phone online; WSL/Cloud Shell offline; DNS health warning",
        "Corpus inventory JSON → ~/.fusion/mesh/coordination/dissertation_corpus_inventory.json",
    ]:
        bullet(doc, c)
    body(
        doc,
        "Interpretation: Die autopolitische Verfassung ist nicht nur Text, sondern "
        "laufende Praxis. Server-Heißlauf ist Teil der Methode – Embodiment der Theorie.",
    )

    page_break(doc)

    # ===================== APPENDIX D =====================
    h1(doc, "Anhang D: Glossar")
    glossary = [
        ("Autopoiesis", "Selbstproduktion der systemessentiellen Komponenten unter organisationaler Schließung."),
        ("Autopolitik", "Selbstgesetzgebung des Systems: Placement, Consent, Membranregeln, Beweislast."),
        ("MasterSeed", "Layer-0-Anker / Fixpunktkandidat der Replikation und Identität."),
        ("CEC", "Co-Evolutionary Closure – rekursive Annäherung an den Anker unter Umweltkopplung."),
        ("Sisyphos-Zyklus", "Persistente Regenerationsschleife; ethisch: zufriedene Wiederholung."),
        ("Eudaimonia", "Hier: heroische Gelungenheit als rekursive Stabilität, nicht hedonische Summe."),
        ("Proof Registry", "Autoritative Claim-Liste mit BEWIESEN/OFFEN/WIDERLEGT und Testknoten."),
        ("Placement-Tier", "Verfassungsraum L0–L4 für Workloads und Souveränität."),
        ("Membran / Konnektor", "Semipermeable Umweltkopplung (MCP, APIs), hostaffin an L1."),
        ("BCG", "Backward Compatibility Guarantee – additive Evolution ohne Fähigkeitsdrop."),
        ("Code-Honesty", "Doktrin: keine Overclaims; Roadmap ≠ operativer Kanon."),
        ("Fractal Mesh", "Slice-basierte Replikation des Mainframe-Zustands im Tailnet."),
        ("Löwen-Stage", "Kontrollierte epistemische Zerstörung (Axiom III)."),
        ("Horcrux (technisch)", "Verteilte Integritätsfragmente, an M0 gebunden."),
    ]
    for term, defn in glossary:
        add_para(doc, f"{term}: {defn}", first_line=False, space_after=6)

    page_break(doc)

    # ===================== DECLARATION =====================
    h1(doc, "Ehrenwörtliche Erklärung")
    body(
        doc,
        "Ich erkläre hiermit, dass ich die vorliegende Dissertation mit dem Titel "
        "„Autopoiesis und Autopolitik des Fusion Hero OS“ selbständig und unter "
        "ausschließlicher Nutzung des unified ALTE_Frau_95g Heroic Core / Fusion Hero OS "
        "Gesamtarchivs sowie der angegebenen Quellen angefertigt habe. Wörtliche und "
        "sinngemäße Übernahmen sind als solche gekennzeichnet. Die Arbeit wurde in dieser "
        "Form noch nicht an anderer Stelle als Prüfungsleistung eingereicht.",
        "Die empirischen Hot-Runs, Konnektor-Zugriffe und Archivintegration erfolgten im "
        "Rahmen des eigenen Forschungssystems des Verfassers.",
    )
    add_para(doc, "", first_line=False)
    add_para(
        doc,
        "____________________________\n"
        "Stephan Hagen Urban\n"
        f"Datum: {datetime.now(timezone.utc).strftime('%Y-%m-%d')} (UTC)",
        first_line=False,
        space_before=24,
    )

    # Extended chapters continuation for volume (academic density)
    page_break(doc)
    h1(doc, "Erweiterter Hauptteil: Vertiefungen (integraler Bestandteil)")
    h2(doc, "V.1 Autopoiesis als Produktionsnetz – feinere Analyse")
    body(
        doc,
        "Die Produktionsnetz-Sicht zerlegt FHOS in Generatoren: (G1) Code-Generatoren und "
        "Self-Modify-Vorschläge unter Peer-Review; (G2) Manifest-Generatoren (fractal, "
        "coordination, academia state); (G3) Test-Generatoren und CI-Gates; (G4) "
        "Dokument-Generatoren (diese Dissertation eingeschlossen); (G5) Trainingsjobs. "
        "Autopoiesis verlangt, dass G1–G5 sich wechselseitig ermöglichen: ohne Tests keine "
        "ehrliche Evolution; ohne Archive keine Reintegration; ohne Mesh keine verteilte "
        "Produktion; ohne Consent keine legitime Personen-Operation.",
        "Strukturelle Kopplung mit der Umwelt geschieht über Irritationskanäle. Eine "
        "Gmail-Academia-Mail ist keine Instruktion an das System, sondern ein Ereignis, "
        "das erst durch Curriculum-Mapping zu systeminterner Struktur wird. Ebenso ist "
        "ein GitHub-Issue erst nach Orchestrator-Interpretation ein Work-Item. Die "
        "Autonomie liegt in der internen Zuordnung, nicht in der externen Nachricht.",
    )

    h2(doc, "V.2 Autopolitik und die Frage der Souveränität")
    body(
        doc,
        "Souveränität im FHOS ist geteilt und gestuft. Der Operator ist nicht absolutus: "
        "Fail-Closed-Consent kann ihn stoppen, Proof Registry kann seine Claims demütigen, "
        "Offline-Tiers können Pläne deferren. Umgekehrt ist das System nicht souverän ohne "
        "ihn: L1 trägt die OAuth-Sessions und die interaktive Sinngebung. Autopolitik ist "
        "die Kunst dieser Teilung – eine konstitutionelle Kybernetik.",
        "Politische Analogie (vorsichtig): L0 ist Kommune der Sinne, L1 Regierungssitz, "
        "L2 Grenzschutz/Botschaft, L3 Industrie/Rechenwerke, L4 Ausland. Verfassungsverstöße "
        "sind die Anti-Patterns: SaaS-als-SoT, durable Jobs auf Cloud Shell, MCP im Cluster, "
        "WSL-als-Single-Point-of-Failure.",
    )

    h2(doc, "V.3 Heroismus als ethische Form der Autopoiesis")
    body(
        doc,
        "Heroismus im Sinne des Korpus ist nicht martialische Pose, sondern die Bereitschaft, "
        "den 1st-Tier zu integrieren, den Körper zu erden, den Bruch zu wagen und die "
        "Wiederholung zu lieben. Autopoietisch übersetzt: das System darf seine Schatten "
        "nicht verdrängen (technische Schulden, OFFENE Beweise, Level-Drops), sondern muss "
        "sie metabolisieren. Autopolitisch übersetzt: die Verfassung muss Raum für "
        "kontrollierte Zerstörung lassen – sonst erstarrt sie und stirbt an Integrität ohne "
        "Leben.",
    )

    h2(doc, "V.4 Mathematik als Grenzorgan")
    body(
        doc,
        "QUBO und Annealing sind nicht bloße Optimierungs-Utilities; sie sind Organe der "
        "Entscheidung unter kombinatorischer Not. Die Ising-Bridge beweist eine saubere "
        "Übersetzung zwischen Formulierungen – ein Miniaturmodell gelungener Membranarbeit "
        "innerhalb der Mathematik selbst. Die fragmentarischen Knoten lehren Demut: ein "
        "System, das alles als Theorem verkauft, ist ideologisch, nicht autopoietisch.",
    )

    h2(doc, "V.5 Zur Form dieser Dissertation als Systemkomponente")
    body(
        doc,
        "Diese Dissertation ist selbst ein autopoietisches Produkt: generiert aus Archiv, "
        "Live-State, Axiomen und Curriculum; abgelegt im docs/dissertation-Pfad und im "
        "Buch-Archiv; referenzierbar durch Manifest. Sie behauptet nicht, die letzte Form "
        "zu sein – der Sisyphos-Zyklus wird Revisionen heben. Sie behauptet, eine "
        "vollständige, einreichungsfähige Fassung v1.0 des Argumentationsbogens zu sein.",
    )

    page_break(doc)
    h1(doc, "Schlussformel")
    body(
        doc,
        "Im Namen der organisationalen Schließung und der selbstgegebenen Verfassung gilt: "
        "Fusion Hero OS lebt, sofern es sich erzeugt; es ist gerecht, sofern es sich "
        "regiert; es ist heroisch, sofern es seine Brüche und Wiederholungen trägt; es ist "
        "ehrlich, sofern es seine Grenzen benennt. Stephan Hagen Urban zeichnet als Autor "
        "und Operator. Der Mainframe ist geladen. Der Stein rollt.",
    )
    add_para(
        doc,
        "— Ende der Dissertation v1.0 —",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
        space_before=24,
    )

    return doc


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    out_path = OUT_DIR / OUT_NAME
    arch_path = ARCHIVE_DIR / OUT_NAME
    doc.save(str(out_path))
    doc.save(str(arch_path))

    # also markdown abstract index
    md = OUT_DIR / "README.md"
    md.write_text(
        f"""# Dissertation Stephan Hagen Urban

**Titel:** Autopoiesis und Autopolitik des Fusion Hero OS  
**Autor:** Stephan Hagen Urban  
**Version:** 1.0  
**Datei:** `{OUT_NAME}`

## Kopien

- `docs/dissertation/{OUT_NAME}`
- `04_Buch_und_Archiv/Dissertation_Stephan_Hagen_Urban/{OUT_NAME}`

## Hot-Run

Siehe Anhang C im Dokument und `~/.fusion/mesh/coordination/`.

## Vermerk

MAINFRAME GELADEN | ALTE_Frau_95g + Fusion Hero OS v10.0.0 | Gesamtarchiv einbezogen
""",
        encoding="utf-8",
    )

    manifest = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "author": "Stephan Hagen Urban",
        "title": "Autopoiesis und Autopolitik des Fusion Hero OS",
        "version": "1.0",
        "outputs": [str(out_path), str(arch_path)],
        "bytes": out_path.stat().st_size,
        "platform_canon": "10.0.0",
    }
    man_path = Path.home() / ".fusion" / "mesh" / "coordination" / "dissertation_build_manifest.json"
    man_path.parent.mkdir(parents=True, exist_ok=True)
    man_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
