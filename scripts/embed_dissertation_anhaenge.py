# -*- coding: utf-8 -*-
"""
Embed docs/dissertation/anhaenge/*.md into the dissertation DOCX
and optionally export PDF via LibreOffice.

Usage:
  python scripts/embed_dissertation_anhaenge.py
  python scripts/embed_dissertation_anhaenge.py --pdf
  python scripts/embed_dissertation_anhaenge.py --regen-catalog --pdf
"""
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import time
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
ANH = ROOT / "docs" / "dissertation" / "anhaenge"
DOCX = (
    ROOT
    / "docs"
    / "dissertation"
    / "Dissertation_Stephan_Hagen_Urban_Autopoiesis_Autopolitik_Fusion_Hero_OS_v1.0.docx"
)
ARCH_DIR = ROOT / "04_Buch_und_Archiv" / "Dissertation_Stephan_Hagen_Urban"
PDF = DOCX.with_suffix(".pdf")
MARKER = "ANHANG_MODULE_FUNKTIONEN_AUS_DEM_NICHTS_V1"

ORDER = [
    "A00_V10_ACTIVATION.md",
    "00_INDEX_ANHAENGE.md",
    "A01_Fundament_aus_dem_Nichts.md",
    "A02_Core_Module_Herleitung.md",
    "A03_Engine_QUBO_Mainframe.md",
    "A04_Methodik_Geltung_V33.md",
    "A05_Meta_Consent_Governance.md",
    "A06_Dashboard_Surfaces_Routes.md",
    "A07_Mesh_Interconnect_Routing.md",
    "A08_Modules_Providers_Orchestration.md",
    "A09_Entwicklungsgeschichte_Null_bis_v10.md",
    "A10_Funktionskatalog_AST.md",
]


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, size=12, bold=False, italic=False, first_line=True, space_after=8):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def h1(doc, t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        set_run_font(r, size=16, bold=True)


def h2(doc, t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        set_run_font(r, size=14, bold=True)


def h3(doc, t):
    p = doc.add_heading(t, level=3)
    for r in p.runs:
        set_run_font(r, size=12, bold=True)


def bullet(doc, t, size=11):
    p = doc.add_paragraph(t, style="List Bullet")
    for r in p.runs:
        set_run_font(r, size=size)


def code_block(doc, text):
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.0
        pf.first_line_indent = Cm(0)
        r = p.add_run(line if line else " ")
        set_run_font(r, name="Consolas", size=9)


def strip_md_inline(s: str) -> str:
    s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
    s = s.replace("**", "").replace("__", "").replace("`", "")
    s = s.replace("*", "")
    return s.strip()


def already_embedded(doc: Document) -> bool:
    for p in doc.paragraphs:
        if MARKER in (p.text or ""):
            return True
    return False


def append_markdown_file(doc: Document, path: Path) -> int:
    """Return approximate word count added."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    words = 0
    in_code = False
    code_buf: list[str] = []
    for line in raw.splitlines():
        if line.strip().startswith("```"):
            if in_code:
                code_block(doc, "\n".join(code_buf))
                words += sum(len(x.split()) for x in code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        if not line.strip():
            continue
        if line.startswith("# "):
            t = strip_md_inline(line[2:])
            h1(doc, t)
            words += len(t.split())
        elif line.startswith("## "):
            t = strip_md_inline(line[3:])
            h2(doc, t)
            words += len(t.split())
        elif line.startswith("### "):
            t = strip_md_inline(line[4:])
            h3(doc, t)
            words += len(t.split())
        elif line.startswith("|") and "|" in line[1:]:
            # tables as monospace lines (portable, no table style dependency)
            code_block(doc, strip_md_inline(line))
            words += len(line.split())
        elif re.match(r"^[-*]\s+", line):
            t = strip_md_inline(re.sub(r"^[-*]\s+", "", line))
            bullet(doc, t)
            words += len(t.split())
        elif re.match(r"^\d+\.\s+", line):
            t = strip_md_inline(re.sub(r"^\d+\.\s+", "", line))
            bullet(doc, t)
            words += len(t.split())
        else:
            t = strip_md_inline(line)
            if t in ("---", "***"):
                continue
            # catalog density: slightly smaller for A10 long lists
            size = 10 if path.name.startswith("A10") else 12
            first = not path.name.startswith("A10")
            add_para(doc, t, size=size, first_line=first, space_after=4 if path.name.startswith("A10") else 8)
            words += len(t.split())
    if in_code and code_buf:
        code_block(doc, "\n".join(code_buf))
        words += sum(len(x.split()) for x in code_buf)
    return words


def find_soffice() -> Path | None:
    candidates = [
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        Path("/usr/bin/soffice"),
        Path("/usr/bin/libreoffice"),
    ]
    for c in candidates:
        if c.exists():
            return c
    return None


def export_pdf(docx: Path, pdf: Path) -> bool:
    soffice = find_soffice()
    if not soffice:
        print("LibreOffice not found — skip PDF")
        return False
    outdir = docx.parent
    cmd = [
        str(soffice),
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(outdir),
        str(docx),
    ]
    print("PDF export:", " ".join(cmd))
    r = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
    if r.returncode != 0:
        print(r.stdout)
        print(r.stderr)
        return False
    # LO names from docx stem
    produced = outdir / (docx.stem + ".pdf")
    if produced.exists() and produced != pdf:
        shutil.copy2(produced, pdf)
    return pdf.exists()


def word_count(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        n += len((p.text or "").split())
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                n += len((cell.text or "").split())
    return n


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--pdf", action="store_true", help="Export PDF via LibreOffice")
    ap.add_argument(
        "--regen-catalog",
        action="store_true",
        help="Regenerate A10 AST catalog (incl. Dashboard)",
    )
    ap.add_argument(
        "--force",
        action="store_true",
        help="Embed even if marker already present (creates duplicate — avoid)",
    )
    ap.add_argument(
        "--activate-v10",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Activate full Fusion Hero OS v10.0.0 before embed (default: on)",
    )
    ap.add_argument(
        "--dashboard-url",
        default=os.environ.get("FUSION_DASHBOARD_URL", "http://127.0.0.1:8000"),
        help="Dashboard base URL for v10 HTTP activation",
    )
    args = ap.parse_args()

    # --- automatic v10 full activation (default) ---
    if args.activate_v10:
        act = ROOT / "scripts" / "activate_v10.py"
        print("=== auto-activate v10.0.0 ===")
        r = subprocess.run(
            [sys.executable, str(act), "--base", args.dashboard_url],
            cwd=str(ROOT),
        )
        if r.returncode not in (0,):
            # version pin hard-fail; still continue catalog if only HTTP soft issues
            # activate_v10 returns 2 on VERSION mismatch
            if r.returncode == 2:
                print("ERROR: v10 platform pin failed")
                return 2
            print(f"activate_v10 exit {r.returncode} — continuing with local pin")

    if args.regen_catalog:
        gen = ROOT / "scripts" / "generate_anhang_katalog.py"
        r = subprocess.run([sys.executable, str(gen)], cwd=str(ROOT))
        if r.returncode != 0:
            return r.returncode

    if not DOCX.exists():
        print("DOCX missing:", DOCX)
        return 1
    if not ANH.is_dir():
        print("Anhaenge missing:", ANH)
        return 1

    doc = Document(str(DOCX))
    if already_embedded(doc) and not args.force:
        print("Anhänge already embedded (marker found). Use --force to re-add.")
        # still allow PDF export
        if args.pdf:
            export_pdf(DOCX, PDF)
        return 0

    doc.add_page_break()
    h1(doc, "Anhang: Module und Funktionen — Herleitung aus dem Nichts")
    add_para(
        doc,
        "Die folgenden Anhänge A01–A10 gehören integral zur Dissertation. "
        "Sie leiten die Module und Funktionen des Fusion Hero OS aus dem Nichts her "
        "und folgen der Designvorlage Kompendium der Heroik V3.3 (Synthese, sechs Bögen, "
        "Geltungsmarken Satz/Bedingt/Modell/Fragment, Trennung von Spezifikation, "
        "Heroischem Exkurs und Herleitung). Der Text ist eine Ausdrucksform; "
        "das laufende OS ist die Dissertation.",
        first_line=True,
    )
    add_para(
        doc,
        f"[{MARKER}] Eingebettet: {datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%MZ')}",
        first_line=False,
        italic=True,
        size=10,
    )

    total_w = 0
    embedded = []
    for name in ORDER:
        path = ANH / name
        if not path.exists():
            print("skip missing", name)
            continue
        doc.add_page_break()
        w = append_markdown_file(doc, path)
        total_w += w
        embedded.append({"file": name, "approx_words": w})
        print(f"  + {name} (~{w} words)")

    # theses for ontology
    doc.add_page_break()
    h1(doc, "Anhang-Nachwort: Dissertation-as-OS und V3.3")
    add_para(
        doc,
        "Dissertation ≡ Fusion Hero OS. Der Text (einschließlich dieser Anhänge) ist "
        "eine Form des Ausdrucks neben Code, Mesh, Dashboard, Tests und Proof Registry. "
        "Arbeitsqualität darf nicht hinter V3.3 zurückfallen: keine Metapher als Beweis, "
        "Geltungsmarken an zentralen Claims, BCG für ältere Kerntexte.",
    )

    doc.save(str(DOCX))
    ARCH_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, ARCH_DIR / DOCX.name)

    wc = word_count(doc)
    man = {
        "embedded_at": datetime.now(timezone.utc).isoformat(),
        "marker": MARKER,
        "docx": str(DOCX),
        "archive_copy": str(ARCH_DIR / DOCX.name),
        "files": embedded,
        "approx_words_added": total_w,
        "approx_words_total": wc,
        "pdf": None,
    }

    pdf_ok = False
    if args.pdf:
        pdf_ok = export_pdf(DOCX, PDF)
        man["pdf"] = str(PDF) if pdf_ok else None
        if pdf_ok:
            shutil.copy2(PDF, ARCH_DIR / PDF.name)

    man_path = Path.home() / ".fusion" / "mesh" / "coordination" / "dissertation_anhaenge_embed.json"
    man_path.parent.mkdir(parents=True, exist_ok=True)
    man_path.write_text(json.dumps(man, indent=2, ensure_ascii=False), encoding="utf-8")
    print("saved", DOCX, "words≈", wc, "added≈", total_w, "pdf", pdf_ok)
    print("manifest", man_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
