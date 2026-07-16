#!/usr/bin/env python3
"""
Build Geisteskrankheiten 4D-Matrix Kompendium v10.0.0

BCG (Backward Compatibility Guarantee):
  - Entire v7 Kompendium body is preserved verbatim after the v10 front matter.
  - New material is strictly additive (Parts VII–X + v10 Synthese-Erweiterung).

Outputs:
  docs/kompendium/geisteskrankheiten-4d/
  03_Code/core/knowledge/
  Desktop + Drive Streaming copies
"""
from __future__ import annotations

import json
import shutil
import sys
import time
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
KNOWLEDGE = ROOT / "03_Code" / "core" / "knowledge"
SCRIPTS = ROOT / "03_Code" / "Dashboard" / "scripts"
OUT_DIR = ROOT / "docs" / "kompendium" / "geisteskrankheiten-4d"
V7 = KNOWLEDGE / "Geisteskrankheiten_4D_Matrix_v7_Kompendium.md"
VERSION = "10.0.0"
TAG = "geisteskrankheiten-4d-v10.0.0"

V10_FRONT = r'''# Geisteskrankheiten in der 4D-Matrix
### Einordnung nach Ursachdimension — MER · KLINIK · RAUM · PHILOSOPHIE · FUSION HERO OS

> **Version 10.0.0** · additive Evolution über **v7 Kompendium** (vollständig erhalten)  
> Duktus: Kompendium der Heroik V3.3 · Plattform: Fusion Hero OS **v10.0.0**  
> Autor/Operator: Stephan Hagen Urban · Stand: ''' + datetime.now(timezone.utc).strftime("%Y-%m-%d") + r'''

---

## BCG — Backward Compatibility Guarantee

Dieses Werk ist **streng rückwärtskompatibel** zur v7-Kompendium-Fassung:

1. **Gesamter v7-Text** (Synthese, §§0–29, 24-Störungs-Matrix, Raum, Trajektorien, Dissertation-/Edition-Brücken) bleibt **wörtlich erhalten** im Abschnitt *»v7-Kern (erhalten)«*.
2. **Neue Inhalte** stehen ausschließlich in der **v10-Synthese-Erweiterung** und in den **Teilen VII–X**.
3. Keine Störung, keine Formel und keine ethische Grenze aus v7 wird stillschweigend gelöscht oder umbenannt.
4. Geltungsmarken bleiben: **Satz · Bedingt · Modell · Fragment** (wie Heroik-Kompendium & Code-Honesty / Proof Registry).

---

## v10-Synthese-Erweiterung: Vier Organe, eine Verfassung

v7 kannte drei Organe: klinische Matrix, formale Brücke, philosophische Einbettung.  
v10 fügt ein **viertes Organ** hinzu, ohne die drei zu ersetzen:

| Organ | Aufgabe | v7 | v10 |
|---|---|---|---|
| **Klinik** | Verorten Z=(K,G,S,N)+I(Z) | ✓ | ✓ erhalten |
| **Formal** | ψ▷φ, Layer, QUBO, DMN | ✓ | ✓ + Proof Registry |
| **Philosophie** | Hospitalismus, Rekonstruktivismus, Schamgrenze | ✓ | ✓ + Axiome I–IV |
| **Autopolitik** | Consent, Placement, Membran, PII | — | **neu** |

Die klinische Landkarte ohne Autopolitik riskiert, dass sensible Zustandsdaten in SaaS-Membranen oder Cluster-Logs landen, ohne Fail-Closed-Consent. Autopolitik ohne Klinik ist leere Governance. v10 koppelt beides.

**Oberste Direktive (klinisch gelesen, Modell):**  
Erhalte und steigere die *rekonstruktive Kohärenz* und *autobiographische Kontinuität* der Person in der nicht-kommutativen Verkettung von Fluss und Schnitt (q∘b) — durch kontrollierte Nothing-Bereitschaft, memetische (nicht mimetische) Heilungspfade und **fail-closed** Personenbezug.

---

## Was v10 zusätzlich leistet (ohne v7 zu ersetzen)

- Mapping **Heroismus-Axiome I–IV** → 4D-Achsen und I(Z)
- **Proof-Honesty:** was in FHOS *BEWIESEN / OFFEN* ist vs. was in der Matrix *Modell* bleibt
- **Consent / Stage-A/B / PII:** klinische Daten nie unkontrolliert in MCP/L3
- **Placement L0–L4** für Therapie-Artefakte (Phone-Journal, Mainframe, Mesh, Cluster)
- **Harmonisierung & Geisterjagd** (AscensionOS loadable) als optionale α₂-Instrumente
- **Sisyphos / CEC / Persistent cycle** als Rückfall- und Regenerationsmetrik
- **Crosswalk** zur Dissertation *Autopoiesis und Autopolitik des Fusion Hero OS*
- Erweiterte **Anti-Patterns** (SaaS-als-SoT für Klinik, Diagnose-by-LLM ohne Hygiene)

## Was v10 weiterhin *nicht* leistet

- Kein DSM/ICD-Ersatz
- Keine Diagnose aus Chat allein
- Kein Wirksamkeitsbeweis durch Landkarte
- Keine stillschweigende Überhöhung von OFFENEN Math-Knoten zu klinischen Gesetzen

---

## Architektur v10 (erweitert)

| Teil | Inhalt | Herkunft |
|---|---|---|
| **0** | v10-Synthese, BCG, Direktive | **v10 neu** |
| **I–VI** | §§0–29 MER, 24 Störungen, Raum, Trajektorien, Dissertation, Edition | **v7 erhalten** |
| **VII** | Axiome · Autopoiesis · Autopolitik der Klinik | **v10 neu** |
| **VIII** | Infrastruktur: Mesh, Consent, Placement, Connectors | **v10 neu** |
| **IX** | Operator-Stack v10 + Proof-Honesty-Tabelle | **v10 neu** |
| **X** | Glossar-Erweiterung, Checklisten, Publikationsvermerk | **v10 neu** |

---

'''

V10_EXTENSION = r'''

---

# TEIL VII — Heroische Axiome, Autopoiesis und Autopolitik der Klinik (v10)

## 30 — Axiome I–IV auf die 4D-Matrix abgebildet

| Axiom | Leitsatz (kurz) | 4D-Bezug | Klinische Lesart (Modell) |
|---|---|---|---|
| **I 1st-Tier** | Eudaimonia beginnt im Schlamm | N/K roh, S-Scham | Integration statt Verleugnung des „Teufels“; N-primäre Last nicht moralisieren |
| **II Somatic** | Körper als Hardware-Horkrux | **K**-Achse, Interozeption | Embodiment vor reiner G-Intervention; Panik/Somatoform ehrlich am Körper |
| **III Psycholyse** | Kontrollierter Bruch / Löwen-Stage | α₂, I(Z)-Senkung | Nur mit Indikation/Kontraindikation (Kap. 18); nie als unkontrollierte Zerstörung |
| **IV CEC / Sisyphos** | Rekursive Closure, zufriedene Wiederholung | Trajektorien Z(t), β-Rückfall | Rückfall = Rekalibrierung möglich; Persistenz der Hebe-Bewegung |

**Geltung:** Mapping = **Modell**, nicht Satz. Keine Axiom-Ableitung ersetzt Leitlinien.

### 30.1 I(Z) und Axiom I

Hohes I(Z) bei N-primären Störungen, die als „Charakterfehler“ gelesen werden, ist ein *Scham-Overlay* (Δ_N). Axiom I verlangt: die biologische Last anerkennen, statt sie moralisch zu vernichten — sonst steigt I(Z).

### 30.2 Axiom II und K-primäre Pfade

Panik und Somatoform: der Körper ist Zeuge. v10-Regel: α₁-Stabilisierung **K** vor reiner G-Umdeutung, wenn autonomes System dominiert (vgl. v7 §3.4, §9).

### 30.3 Axiom III und memetische vs. mimetische Heilung

Psycholyse/Bruch ohne Stamm und ohne Hygiene erzeugt **mimetische** Scheinheilung (Symptomkosmetik) oder iatrogene Fragmentierung. v7 §6 bleibt bindend; v10 bindet Axiom III an §18.5 Kontraindikationen.

### 30.4 Axiom IV und Trajektorien

Limit-Cycle (Bipolar, OCD) und Trauma-Attraktor (PTBS) sind Sisyphos-Strukturen: der Stein rollt. α-Phasen sind Hebeversuche; β ist der Abstieg. CEC-Lesen: jede Runde soll d_I zur rekonstruktiven Kohärenz *kontraktiv* sein — **Modell**, empirisch zu prüfen.

---

## 31 — Autopoiesis der klinischen Erkenntnis

Ein klinisches Wissenssystem ist autopoietisch, wenn es:

1. seine **Begriffe und Profile** rekursiv erzeugt und prüft (Matrix-Updates, Review),
2. eine **Membran** hat (Consent, Anonymisierung, keine unkontrollierte SaaS-SoT),
3. einen **Anker** hat (klinische Hygiene §0, Leitlinien-Vorrang bei Gefahr),
4. **regeneriert** (Edition History, Sisyphos der Revision, Proof-Honesty).

Die 4D-Matrix ist *Werkzeugorgan*, nicht Souverän über den Menschen. Der Mensch bleibt Source-of-Truth seiner Einwilligung.

---

## 32 — Autopolitik: wer darf klinische Zustände wohin legen?

| Placement | Erlaubt für | Verboten für |
|---|---|---|
| **L0 Phone** | Journal, Stimmung, private Capture | ungefilterte Diagnose-Labels an Dritte |
| **L1 Mainframe** | lokale Auswertung, Orchestrator, MCP *mit Consent* | stilles Training mit PII |
| **L2 Mesh-Exit** | verschlüsselte Replica ohne Klartext-PII | Klartext-Psychprofile |
| **L3 Cluster** | anonymisierte Batch-Stats, synthetische Szenarien | identifizierende Kliniktagebücher |
| **L4 SaaS** | nur membranisiert, minimiert, widerrufbar | Source-of-Truth der Identität/Klinik |

**Anti-Patterns (v10):**

- Diagnose-by-LLM ohne Hygiene und ohne Mensch-in-the-Loop  
- SaaS als alleinige Akte  
- Durable Klinik-Jobs auf Cloud-Shell-Nodes  
- MCP-Export von Rohjournalen ohne Stage-A Consent  

---

# TEIL VIII — Infrastruktur-Brücke Fusion Hero OS v10 (v10)

## 33 — Module und Dateien (ehrliche Landkarte)

| Modul / Artefakt | Pfad / Ort | Bezug 4D | Status |
|---|---|---|---|
| 4D-Kompendium v7 | `03_Code/core/knowledge/…_v7_Kompendium.md` | Kern | erhalten |
| HarmonisierungsCore | `ascension_os/core/harmonisierung_module.py` | α₂, q∘b | loadable |
| Geisterjagd | `ascension_os/core/geisterjagd_module.py` | Ghost Fixing §16.5 | loadable |
| Consent gate | `ascension_os/consent_gate.py` | Autopolitik | operativ v10 |
| Proof Registry | `proof_registry.yaml` | Epistemik §17.6 | operativ |
| Mesh coordination | `mesh_service_coordination.yaml` | Placement | operativ |
| Dashboard | `03_Code/Dashboard` | Operator-UI | operativ |
| Exposure practice | `exposure_practice_module` | S-Exposition *ohne* echte Dritte | v9.7 |

## 34 — Consent & PII (Stage-A/B)

v10 Stage-A/B: fail-closed für personenbezogene Ops, PII-Cleanup, Persona-Scanner.  
**Klinische Anwendung:** Jede Speicherung von Z-Profilen mit Bezug zu einer realen Person setzt expliziten Consent voraus. Default = **deny**.

## 35 — QUBO-Triage und Cluster

v7 §17.4 bleibt: welche Achse zuerst.  
v10: schwere Annealing-/Batch-Jobs → **L3**, interaktive Triage-UI → **L1**. Ergebnisse fließen als *aggregierte* Hinweise zurück, nicht als Identitätsersatz.

---

# TEIL IX — Operator-Stack v10 und Proof-Honesty (v10)

## 36 — Erweiterte Gesamtpipeline (v7 §28 + v10)

```
[Gefahr?] --ja--> Standardversorgung (DSM/ICD/Notfall)
    |
   nein
    v
[Consent?] --nein--> STOP (fail-closed)
    |
   ja
    v
Z-Erhebung (K,G,S,N) + I(Z)-Triangulation
    v
Cluster (N/S/G/K-primär) + Komorbidität (§7)
    v
α₁ Stabilisierung → α₂ Rekonstruktion → α₃ Integration
    v
Memetisch? (Stamm, Embodiment, Beziehung) vs mimetisch (Kosmetik)
    v
Optional: Harmonisierung / Geisterjagd / DMN-Protokoll (§18) — nur indiziert
    v
Trajektorie Z(t) dokumentieren (Sisyphos-Log)
    v
Edition / Matrix-Revision (BCG, kein Overclaim)
```

## 37 — Proof-Honesty-Kreuztabelle (Auswahl)

| Claim-Bereich | In Matrix | In FHOS Proof Registry | Lesart |
|---|---|---|---|
| Z=(K,G,S,N) Landkarte | Modell | — | Werkzeug |
| I(Z)=Σ\|Δi\| | Modell / operationalisierbar | — | Approximation |
| q∘b nicht-kommutativ (Math) | Brücke | Kommutator-Antisymmetrie BEWIESEN (Algebra) | nicht = klinischer Beweis |
| Banach-Fixpunkt universell | — | OFFEN | nicht klinisch behaupten |
| parallel_anneal kleine Instanzen | QUBO-Triage-Idee | BEWIESEN (eng) | nicht Therapie-Wirksamkeit |
| Mutual sync / Semilattice | — | BEWIESEN (eng) | Mesh, nicht Psyche |
| Consent fail-closed | Autopolitik v10 | Stage-A Tests | operativ |

**Regel:** Matrix-Modelle dürfen FHOS-Beweise *inspirieren*, nicht *usurpieren*.

## 38 — Störungen × v10-Werkzeug (Erweiterung zu §19.3)

| Störung (Auswahl) | v7-Kern | v10-Zusatz |
|---|---|---|
| Depression endogen | N→K, α₁ med + Stamm | Axiom I gegen Moralisierung; Sisyphos-Log der Hebe-Tage |
| Borderline | S-Sprungfeld | Consent-strikte Journalführung L0/L1; kein LLM-Diagnose-Dump |
| PTBS | Trauma-Attraktor | Geolock/Trigger-Hygiene; optional exposure_practice *simuliert* |
| OCD | CSTC, b-Übersteuerung | QUBO-Triage: G/K-Schleife zuerst; Limit-Cycle tracken |
| Narzissmus | G+ i grandios | Geisterjagd auf Grandiositäts-Fixierung = **Modell**; Peer-Review-Ethik |
| Autismus Masking | I hoch bei ● nahe Z* | Axiom II Erschöpfungs-K; keine „Reparatur“ der Natur-Identität |

---

# TEIL X — Checklisten, Glossar-Erweiterung, Vermerk (v10)

## 39 — Kliniker-Checkliste (5 Minuten)

1. Akute Gefahr? → Standard, Matrix pausiert.  
2. Consent dokumentiert?  
3. Primärcluster N/S/G/K + Top-2 Komorbiditäten.  
4. I(Z): wo Δ am größten (Selbst vs. Fremd)?  
5. α-Phase jetzt? memetisch oder mimetisch?  
6. Daten: welches Placement? PII minimiert?  
7. Overclaim-Check: Satz/Bedingt/Modell/Fragment gesetzt?

## 40 — Glossar-Erweiterung v10

| Term | Kurz |
|---|---|
| **Autopolitik (klinisch)** | Selbstgesetzgebung über Placement, Consent, Membran für Zustandsdaten |
| **BCG** | Backward Compatibility — v7 bleibt erhalten |
| **CEC** | Co-Evolutionary Closure — kontraktive Regeneration (Modell) |
| **Fail-closed** | Ohne Consent keine personenbezogene Op |
| **L0–L4** | Placement-Tiers Edge→SaaS |
| **Memetisch/Mimetisch** | Echte vs. kosmetische Heilung (v7 §6) |
| **SoT** | Source of Truth — Klinikidentität nie allein in L4 |

## 41 — Publikations- und Editionsvermerk v10.0.0

- **Plattform:** Fusion Hero OS v10.0.0  
- **Vorgänger:** Geisteskrankheiten_4D_Matrix **v7** Kompendium (vollständig eingebettet)  
- **Stilreferenz:** Kompendium der Heroik V3.3  
- **Verwandt:** Dissertation *Autopoiesis und Autopolitik des Fusion Hero OS* (Release `dissertation-v1.0`)  
- **Ehrlichkeit:** Dieses Dokument ist ein **konzeptuelles und didaktisches** Werk. Es ersetzt keine ärztliche oder psychotherapeutische Leistung.

---

## Schlusswort v10

Die Landkarte wurde nicht neu erfunden — sie wurde **verfasst**.  
v7 gab den Raum; v10 gibt der Klinik eine Verfassung für Daten, Consent und epistemische Demut.  
Der Stein der Revision wird erneut gehoben. Das ist kein Scheitern. Das ist die Form.

— Ende der v10-Erweiterungen; es folgt der vollständige v7-Kern (BCG) —

'''


def _strip_duplicate_v7_header(body: str) -> str:
    """Keep full v7 body; only drop redundant top title block if identical."""
    text = body
    # If starts with same title as front matter, keep everything after first architecture table
    # Actually BCG says preserve entire v7 - include full file as-is under a banner.
    return text


def build_markdown() -> str:
    if not V7.is_file():
        raise FileNotFoundError(f"Missing v7 source: {V7}")
    v7 = V7.read_text(encoding="utf-8")
    banner = (
        "\n\n---\n\n"
        "# v7-Kern (erhalten — BCG)\n\n"
        "> Ab hier beginnt der **vollständige** Text der v7-Kompendium-Fassung "
        "(unverkürzt). Alle Abschnitte 0–29 bleiben kanonischer Bestandteil von v10.0.0.\n\n"
        "---\n\n"
    )
    return V10_FRONT + V10_EXTENSION + banner + _strip_duplicate_v7_header(v7)


def build_docx(md_text: str, path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    def font(run, size=11, bold=False, italic=False):
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    # Cap body length in DOCX for Word practicality: full MD is source of truth;
    # DOCX gets front + extension + condensed note pointing to full MD/PDF.
    # User asked for full content - include as much as feasible.
    # 91k chars is fine for docx.
    for line in md_text.splitlines():
        s = line.rstrip()
        if not s:
            doc.add_paragraph()
            continue
        if s.startswith("# "):
            p = doc.add_heading(s[2:].strip(), level=1)
            for r in p.runs:
                font(r, 16, bold=True)
        elif s.startswith("## "):
            p = doc.add_heading(s[3:].strip(), level=2)
            for r in p.runs:
                font(r, 13, bold=True)
        elif s.startswith("### "):
            p = doc.add_heading(s[4:].strip(), level=3)
            for r in p.runs:
                font(r, 12, bold=True)
        elif s.startswith("> "):
            p = doc.add_paragraph()
            r = p.add_run(s[2:])
            font(r, 10, italic=True)
        elif s.startswith("|") and s.endswith("|"):
            p = doc.add_paragraph()
            r = p.add_run(s)
            font(r, 8)
            p.paragraph_format.space_after = Pt(0)
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(s[2:], style="List Bullet")
            for r in p.runs:
                font(r, 11)
        elif s.startswith("```"):
            continue
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            r = p.add_run(s)
            font(r, 11)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_pdf(md_text: str, path: Path) -> None:
    sys.path.insert(0, str(SCRIPTS))
    try:
        from kompendium_pdf_renderer import render_kompendium_pdf

        # Patch version globals if present
        import kompendium_pdf_renderer as kpr

        kpr.DOC_SHORT = "Geisteskrankheiten in der 4D-Matrix"
        kpr.DOC_VERSION = "v10.0.0"
        kpr.DOC_EDITION = "BCG v7 + Fusion Hero OS · " + datetime.now(timezone.utc).strftime("%Y-%m")
        render_kompendium_pdf(
            md_text,
            path,
            title="Geisteskrankheiten in der 4D-Matrix",
            subtitle="MER · KLINIK · RAUM · PHILOSOPHIE · FUSION HERO OS",
            tagline="v10.0.0 — additive Evolution über das v7-Kompendium (vollständig erhalten)",
            version="v10.0.0 · BCG · mit v7-Kern",
        )
    except Exception as exc:  # noqa: BLE001
        # Fallback: LibreOffice from a minimal HTML or skip
        print("PDF renderer failed:", exc)
        # Try reportlab minimal
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        from reportlab.lib.units import cm
        from html import escape

        doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm)
        styles = getSampleStyleSheet()
        story = []
        for line in md_text.splitlines()[:4000]:
            if line.startswith("#"):
                story.append(Paragraph(escape(line.lstrip("# ").strip()), styles["Heading1"]))
            else:
                story.append(Paragraph(escape(line[:2000]), styles["Normal"]))
            story.append(Spacer(1, 4))
        doc.build(story)


def copy_desktop_drive(md: Path, pdf: Path, docx: Path) -> dict:
    targets = {}
    desktop = Path.home() / "OneDrive" / "Desktop"
    if not desktop.is_dir():
        desktop = Path.home() / "Desktop"
    desk_dir = desktop / "Forschung"
    desk_dir.mkdir(parents=True, exist_ok=True)
    for src, name in ((md, md.name), (pdf, pdf.name), (docx, docx.name)):
        if src.is_file():
            dst = desk_dir / name
            shutil.copy2(src, dst)
            targets[f"desktop_{name}"] = str(dst)

    gdrive = Path.home() / "Google Drive-Streaming" / "FusionHero_Offload" / "kompendium"
    gdrive.mkdir(parents=True, exist_ok=True)
    for src in (md, pdf, docx):
        if src.is_file():
            dst = gdrive / src.name
            shutil.copy2(src, dst)
            targets[f"gdrive_{src.name}"] = str(dst)
    return targets


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    md_text = build_markdown()
    stem = f"Geisteskrankheiten_4D_Matrix_v{VERSION}_Kompendium"
    md_path = OUT_DIR / f"{stem}.md"
    pdf_path = OUT_DIR / f"{stem}.pdf"
    docx_path = OUT_DIR / f"{stem}.docx"
    md_path.write_text(md_text, encoding="utf-8")
    # knowledge mirror
    (KNOWLEDGE / f"{stem}.md").write_text(md_text, encoding="utf-8")

    build_docx(md_text, docx_path)
    build_pdf(md_text, pdf_path)

    copies = copy_desktop_drive(md_path, pdf_path, docx_path)

    # also knowledge pdf/docx if exist
    if pdf_path.is_file():
        shutil.copy2(pdf_path, KNOWLEDGE / pdf_path.name)
    if docx_path.is_file():
        shutil.copy2(docx_path, KNOWLEDGE / docx_path.name)

    idx = {
        "id": "geisteskrankheiten_4d_v10_kompendium",
        "version": VERSION,
        "tag": TAG,
        "bcg_source": str(V7),
        "updated_ts": time.time(),
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "chars": len(md_text),
        "files": {
            "md": str(md_path),
            "pdf": str(pdf_path) if pdf_path.is_file() else None,
            "docx": str(docx_path) if docx_path.is_file() else None,
        },
        "copies": copies,
        "pdf_bytes": pdf_path.stat().st_size if pdf_path.is_file() else 0,
        "docx_bytes": docx_path.stat().st_size if docx_path.is_file() else 0,
    }
    (OUT_DIR / "geisteskrankheiten_4d_v10_index.json").write_text(
        json.dumps(idx, indent=2), encoding="utf-8"
    )
    (KNOWLEDGE / "geisteskrankheiten_4d_v10_kompendium.json").write_text(
        json.dumps(idx, indent=2), encoding="utf-8"
    )
    pub = OUT_DIR / "PUBLICATION.md"
    pub.write_text(
        f"""# Publication — Geisteskrankheiten 4D Matrix v{VERSION}

**Tag:** `{TAG}`  
**BCG:** full v7 Kompendium embedded  
**Platform:** Fusion Hero OS v10.0.0  

## Files

- `{md_path.name}`
- `{pdf_path.name}`
- `{docx_path.name}`

## Release (after push)

https://github.com/95guknow/fusion-hero-os/releases/tag/{TAG}
""",
        encoding="utf-8",
    )
    print(json.dumps(idx, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
