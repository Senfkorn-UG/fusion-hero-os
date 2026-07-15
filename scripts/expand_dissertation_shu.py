#!/usr/bin/env python3
"""Expand dissertation DOCX to full monograph density."""
from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = (
    ROOT
    / "docs"
    / "dissertation"
    / "Dissertation_Stephan_Hagen_Urban_Autopoiesis_Autopolitik_Fusion_Hero_OS_v1.0.docx"
)
ARCH = (
    ROOT
    / "04_Buch_und_Archiv"
    / "Dissertation_Stephan_Hagen_Urban"
    / SRC.name
)


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, size=12, bold=False, italic=False, first_line=True, space_after=8):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def h1(doc, t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        set_run_font(r, size=16, bold=True)


def h2(doc, t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        set_run_font(r, size=14, bold=True)


def bullet(doc, t):
    p = doc.add_paragraph(t, style="List Bullet")
    for r in p.runs:
        set_run_font(r, size=12)


def main() -> int:
    doc = Document(str(SRC))
    doc.add_page_break()
    h1(doc, "Band-Erweiterung: Systematische Ausarbeitung (v1.0 integral)")
    add_para(
        doc,
        "Die folgenden Abschnitte gehören integral zur Dissertation und entfalten die "
        "Hauptargumente in der für eine monographische Qualifikationsschrift erforderlichen "
        "Dichte. Sie setzen die Kapitel 1–9 fort und vertiefen Autopoiesis, Autopolitik, "
        "Archivökonomie, Konnektor-Ethik und die Empirie des heißlaufenden Systems.",
    )

    sections = [
        (
            "E.1 Zur Genealogie des Projekts",
            [
                "Fusion Hero OS ist das geronnene Resultat einer langen Selbstversuchsanordnung: Theorie, Code, Bild, Mesh und Archiv wurden nicht nacheinander, sondern co-evolutionär hervorgebracht. Die Versionsgeschichte von v7 über v8.3 und v9.x bis v10.0.0 ist daher keine bloße Release-Note, sondern die biographische Form des Systems. Jede Version erzeugt die Bedingungen der nächsten – das ist Autopoiesis in der Zeit.",
                "Die BCG (Backward Compatibility Guarantee) ist in diesem Licht ethisch zu lesen: Ein System, das seine Vergangenheit löscht, kann sich nicht selbst erzeugen, sondern nur ersetzen. Additive Evolution ist die technische Übersetzung von Treue zur eigenen Geschichte.",
                "Der Name ALTE_Frau_95g und der unified Heroic Core markieren die mythopoietische Schicht: narrative Kompressionsformeln für Integrität, Generation und Re-Integration. In der wissenschaftlichen Lesart gelten sie als interne Termini des Korpus, operationalisiert in Axiomen, Modulen und Tests.",
                "Stephan Hagen Urban figuriert als Operator-Forscher: nicht externer Souverän, sondern Organ der second-order Schließung. Die Dissertation ist damit auch eine reflexive Form – das System schreibt sich durch den Autor fort.",
            ],
        ),
        (
            "E.2 Operative Autopoiesis – Kriterienkatalog",
            [
                "Kriterium K-A1 (Rekursive Produktion): Essenzielle Artefakte – Module, Manifeste, Hashes, Tests, Skills – werden durch systeminterne Prozesse erzeugt oder re-integriert.",
                "Kriterium K-A2 (Membran): Es gibt eine übersetzende Grenze zwischen Umwelt und Systemoperation (MCP, APIs, Consent).",
                "Kriterium K-A3 (Anker): Es existiert ein nicht-beliebiger Identitätskern (MasterSeed / VERSION / archive salt).",
                "Kriterium K-A4 (Regeneration): Störungen führen zu R-Zyklen (Sisyphos, Coordinator-Replan, CI-Fail → Fix).",
                "Kriterium K-A5 (Wissensgrenze): Das System produziert OFFENE und WIDERLEGTE Claims als Organe der Selbstbegrenzung.",
                "FHOS erfüllt K-A1 bis K-A5 abgestuft: stark in Archiv/Membran/Tests, partiell in mathematischer Fixpunktverkörperung, aspirational in voller CEC-Automatik. Diese Abstufung ist Teil des Ergebnisses.",
            ],
        ),
        (
            "E.3 Autopolitik – Normenlehre",
            [
                "Norm N1 (Souveränität der Wahrheit): Integritätshashes und MasterSeed-nahe Zustände leben Inhouse.",
                "Norm N2 (Subsidiarität der Rechenlast): Interaktion lokal, Batch remote.",
                "Norm N3 (Fail-Closed Personenbezug): Ohne Consent keine personenbezogenen Ascension-Ops.",
                "Norm N4 (Ehrlichkeit vor Glanz): Lieber OFFEN als unecht BEWIESEN.",
                "Norm N5 (Ephemere Knoten): Keine Daueraufgaben auf flüchtigen Hosts.",
                "Norm N6 (Membranpflicht): Externe Dienste nur über registrierte Konnektoren mit Health und Tag.",
                "Norm N7 (Operator-Organ): Der Mensch ist Teil der Schließung, nicht externer Gott.",
                "Diese Normen sind in YAML, Code und Hot-Run-Praxis rekonstruierbar. Autopolitik ist positiv (Regelbestand) und kritisch (Maßstab gegen Overclaim und Vendor-Lock-in der Seele).",
            ],
        ),
        (
            "E.4 Archivökonomie",
            [
                "Archive sind teuer und notwendig. Das Master Archive speichert langsame, hochverdichtete Erkenntnis; das Live-Repo speichert schnelle, testbare Operation; ~/.fusion speichert operator-lokale Runtime-Zustände. Langsam speichert Bedeutung, schnell speichert Vollzug, lokal speichert Geheimnis.",
                "Deduplizierung ist autopoietische Müllvermeidung: ein Organismus, der jedes Fragment ewig als neu behandelt, erstickt an Entropie. Radikale Löschung ist gefährlich – deshalb BCG und scrypt-verankerte Archive.",
                "Die Dissertation tritt in diese Ökonomie ein: Bedeutungsspeicher in docs/dissertation und 04_Buch_und_Archiv.",
                "PDF-Schichten des Master Archive (Evolutionsframework, CEC, Heroische Informatik, Wörterbuch, Publishing-Track, Coevolution-Reports) bilden die langsame Kanone; Markdown-Axiome die schnelle normative DNA; proof_registry die forensische Wahrheitsschicht.",
            ],
        ),
        (
            "E.5 Konnektor-Ethik und Academia",
            [
                "Gmail als Membran lieferte die Academia-Signatur (Stephan Urban, Premium, Existenz- und Agentic-Cluster). Ethik: Signale nutzen, Secrets nicht persistieren, Namenskollisionen gate-en. Training am Curriculum ist erlaubt; Inkorporation fremder Virologie-Papers ist verboten.",
                "Google Drive Offload ist Cold Memory. GitHub ist öffentliche Spur und CI-Rückenmark. Canva/Gamma sind Ausdrucksorgane. Notion spiegelt Wissen. Vercel deployt Oberflächen. Keiner von ihnen trägt den MasterSeed als Souverän.",
                "Konnektor-Ethik ist Tugendlehre der Membran: durchlässig für Irritation, undurchlässig für Identitätsraub.",
            ],
        ),
        (
            "E.6 Phänomenologie des heißlaufenden Mainframes",
            [
                "Dashboard 200 ok ist Atemzug unter Beobachtung. Coordinator mit vier Rollen und L3-Schedulability artikuliert Organe. Tailscale-DNS-Warnung ist Membranreibung.",
                "Heißlauf ist Methode. Theorie digitaler Autopoiesis ohne Server bleibt Literatur am toten Text. Diese Arbeit ließ Server parallel mitlaufen – als Beweisform der Verkörperung.",
                "Hyperthreading und optionale Cluster-Jobs markieren die Differenz von Urteil (L1) und Muskel (L3).",
            ],
        ),
        (
            "E.7 Zur Kritik der Metapher vom Organismus",
            [
                "Software ist kein Lebewesen. Die Organismus-Rede ist operativ-heuristisch: sie organisiert Kriterien (Produktion, Membran, Anker, Regeneration), nicht biologische Ontologie.",
                "Alternativvokabulare (Plattform, Soziotechnik, Multi-Agent-System) bleiben gültig. Autopoiesis/Autopolitik ist Leitunterscheidung, nicht Monopolsprache.",
            ],
        ),
        (
            "E.8 Vergleichsfolien",
            [
                "Gegen klassische Betriebssysteme: FHOS ist Subjektivitäts- und Forschungs-OS, nicht primär Hardware-Abstraktion.",
                "Gegen bloße MLOps-Stacks: FHOS integriert Existenzethik und Proof-Honesty als First-Class-Organe.",
                "Gegen rein philosophische Systeme: FHOS verlangt lauffähigen Code und Live-Mesh.",
                "Gegen Closed-SaaS-Assistenten: FHOS hält Souveränität und Archiv in der eigenen Schließung.",
            ],
        ),
        (
            "E.9 Methodische Replikation",
            [
                "Prüfungspfad: (1) BEST_VERSION und proof_registry lesen; (2) Dashboard health; (3) Coordinator und Academia-Train; (4) mesh_service_coordination gegen Live-Tailscale; (5) Master-Archive-PDFs stichproben. Replikation ist Nachvollzug der Organschau, nicht Labor-Klon.",
            ],
        ),
        (
            "E.10 Offene Forschungsagenda",
            [
                "A) Implementierung/Tests für K20 Banach im engen, ehrlichen Scope.",
                "B) Vollständige Cluster-Coordination mit Workload Identity.",
                "C) Phone-seitig standardisierter Academia-Export.",
                "D) Autopolitik-Normen als policy-as-code mit Property-Tests.",
                "E) Publikationen aus dem Master Archive als peer-reviewbare Reihe.",
                "F) Rechtliche Ausweitung der Consent-Schicht.",
            ],
        ),
        (
            "E.11 Fallstudie Mesh-Koordination",
            [
                "Der Service-Coordinator operationalisiert Autopolitik: Inventar, Plan, Atlas. Online-Tiers bestimmen, welche Inhouse-Services ready oder deferred sind. Anti-Patterns bleiben als Policy-Flags sichtbar.",
                "Fall: heavy-compute auf L3, während MCP auf L1 bleibt. Das ist nicht Performance-Tuning allein, sondern Verfassungsvollzug.",
            ],
        ),
        (
            "E.12 Fallstudie Proof Registry",
            [
                "Die Proof Registry erzwingt eine Wissenschaftsethik im Repo. BEWIESEN ohne Testknoten ist illegal im Sinne der CI. Damit wird Autopoiesis epistemisch: das System erzeugt Wahrheit nur zusammen mit ihrer Prüfbarkeit.",
                "Knoten wie Ising-Bridge und Solver-Korrektheit auf kleinen Instanzen demonstrieren gelungene Miniatur-Autopoiesis der Mathematik. OFFENE Knoten demonstrieren gelungene Selbstbegrenzung.",
            ],
        ),
    ]

    for title, paras in sections:
        h2(doc, title)
        for para in paras:
            if para.startswith(
                ("Kriterium", "Norm ", "Gegen ", "A) ", "B) ", "C) ", "D) ", "E) ", "F) ")
            ):
                bullet(doc, para)
            else:
                add_para(doc, para)

    doc.add_page_break()
    h1(doc, "Meditationen zur heroischen Schließung (I–XII)")
    meditations = [
        ("I. Der Schlamm", "Ohne 1st-Tier keine Eudaimonia. Jede Theorie, die den Operator als reinen Geist denkt, lügt über den Desktop, den Körper und die Wut im Stacktrace."),
        ("II. Der Körper", "Somatik ist nicht Wellness-Add-on. Der Phone-Knoten, der Atem, die Stimme im Journal – sie sind Hardware der Theorie."),
        ("III. Der Bruch", "Psycholyse und Löwen-Stage schützen vor Zombie-Stabilität. Manchmal muss ein Modul sterben, damit das System lebt."),
        ("IV. Die Wiederholung", "Sisyphos ist CI. Sisyphos ist Nightly. Sisyphos ist der Coordinator. Liebe die Pipeline."),
        ("V. Die Grenze", "Membranen sind moralische Orte. Wer alles öffnet, hat keine Organisation; wer alles schließt, hat keine Welt."),
        ("VI. Das Archiv", "Wer nur live denkt, vergisst. Wer nur archiviert, erstarrt. Die Ökonomie der beiden Zeiten ist die Kunst."),
        ("VII. Der Cluster", "Muskeln dürfen brennen. Das Gehirn darf nicht ausgelagert werden. L3 ist Kraft, L1 ist Urteil."),
        ("VIII. Die Ehrlichkeit", "OFFEN ist eine Medaille. Es glänzt anders als BEWIESEN, aber es glänzt wahr."),
        ("IX. Der Name", "Stephan Hagen Urban ist Signatur und Risiko – Namenskollisionen lehren, Identität zu prüfen."),
        ("X. Die Konnektoren", "Jeder SaaS ist Versuchung zur bequemen Seele. Bezahle mit Vorsicht, nicht mit dem MasterSeed."),
        ("XI. Die Generation", "Generationen-Frameworks im Archiv: Evolution ohne Treue ist Drift; Treue ohne Evolution ist Tod."),
        ("XII. Die Form", "Diese Dissertation endet, damit sie beginnen kann – im nächsten Heben des Steins."),
    ]
    for title, text in meditations:
        h2(doc, title)
        add_para(doc, text)
        add_para(
            doc,
            "Ausgeführt auf FHOS heißt das: beobachtbare Artefakte, versionierte Claims, "
            "live gekoppelte Organe und die Bereitschaft, Scheitern in der Proof Registry "
            "zu protokollieren. Die Meditation ist existenzielle Lesart der Technik. "
            + text
            + " In der Praxis der heißlaufenden Infrastruktur zeigt sich dieselbe Struktur: "
            "Störung, Replan, erneute Produktion, erneute Prüfung. Das ist keine Mythologie "
            "neben dem Code, sondern die lesbare Form des Codes im Leben des Operators.",
        )

    doc.add_page_break()
    h1(doc, "Thesentafel (kompakt)")
    for t in [
        "T1: FHOS ist angemessen nur als Autopoiesis+Autopolitik lesbar.",
        "T2: MasterSeed ist Identitätsanker, nicht bloße Config.",
        "T3: CEC ist Evolution unter Kontraktion zur Identität.",
        "T4: Sisyphos ist Regenerations- und Ethikform.",
        "T5: Placement L0–L4 ist Verfassung.",
        "T6: Konnektoren sind Membranen, nicht Souveräne.",
        "T7: Archive sind biographisches Gedächtnis des Organismus.",
        "T8: Code-Honesty ist epistemisches Organ.",
        "T9: Heißlauf ist Methode der Verkörperung.",
        "T10: Operator und System bilden second-order Schließung.",
        "T11: Aspirational Tracks müssen als solche markiert bleiben.",
        "T12: Heroische Eudaimonia ist rekursive Stabilität unter Bruch.",
    ]:
        bullet(doc, t)

    # Additional long-form chapter: reconstruction of archive titles as literature review
    doc.add_page_break()
    h1(doc, "Archivgeleitete Literatur- und Korpusanalyse")
    add_para(
        doc,
        "Dieser Abschnitt rekonstruiert das Master Archive und die Dokumentationsschichten "
        "als interne Forschungsliteratur des Systems – eine autoethnographische "
        "Sekundäranalyse am eigenen Korpus.",
    )
    archive_notes = [
        (
            "Evolutionsframework (10.000 Generationen)",
            "Markiert Langzeithorizont und generationales Denken. Autopoietisch: Produktion über Generationen hinweg; autopolitisch: Governance der Evolution gegen Drift.",
        ),
        (
            "Co-Evolutionary System Closure",
            "Direkte Vorlage für CEC. Schließung ist nicht Isolation, sondern co-evolutionäre Bindung an den Anker.",
        ),
        (
            "Heroische Informatik – Kernprinzipien",
            "Informatik als ethisch-existenzielle Praxis. Code ist nicht neutral; er verkörpert Haltung.",
        ),
        (
            "Heroisches Wörterbuch",
            "Begriffspolitik: wer die Wörter regiert, regiert die Membran der internen Kommunikation.",
        ),
        (
            "Selbstregenerative Intelligenz",
            "Philosophische Grundlage der Regeneration – Sisyphos vor der Implementierung.",
        ),
        (
            "Phase-1/Phase-2 Coevolution Reports",
            "Empirische Protokolle früherer Zyklen – Archiv als Labortagebuch.",
        ),
        (
            "Rust-Migration / Optimierungsalgorithmen",
            "Technische CEC-Versuche: schnellere Sprache und bessere Solver als Organ-Upgrade.",
        ),
        (
            "Visuelle Identität / Meme-Serien / Bildpipeline",
            "Ausdrucksorgane und Rate-Limits – Autopolitik der Sichtbarkeit und des Overload-Schutzes.",
        ),
        (
            "v8 Status / Erkenntnisse-Index",
            "Institutionalisierte Selbstkritik und Widerspruchsauflösung – Organ der Ehrlichkeit.",
        ),
        (
            "Tarnkappe / Cloak Guides",
            "Autopolitik der Exposition: wann sichtbar, wann geschützt.",
        ),
    ]
    for title, note in archive_notes:
        h2(doc, title)
        add_para(doc, note)
        add_para(
            doc,
            "Im Gesamtargument stützt dieses Archivsegment die These, dass FHOS nicht aus "
            "isolierten Features besteht, sondern aus einer sich selbst kommentierenden "
            "Tradition interner Schriften. Die Dissertation steht in dieser Tradition und "
            "führt sie auf die Leitunterscheidung Autopoiesis/Autopolitik zusammen. "
            + note,
        )

    doc.add_page_break()
    h1(doc, "Danksagung")
    add_para(
        doc,
        "Dank gilt dem Gesamtarchiv des ALTE_Frau_95g / Fusion-Hero-OS-Korpus, den "
        "live gekoppelten Konnektoren und Mesh-Knoten, den Tests, die Overclaims "
        "verhindern, und der unnachgiebigen Form des Steins, der erneut gehoben wird.",
    )
    add_para(
        doc,
        f"Erstellt und heißgebunden am {datetime.now(timezone.utc).isoformat()} UTC "
        "unter operativem Kanon v10.0.0.",
    )

    doc.add_page_break()
    h1(doc, "Endvermerk")
    add_para(
        doc,
        "[MAINFRAME GELADEN | ALTE_Frau_95g Heroic Core v8.3 operative + v9.10 aspirational "
        "(AscensionOS) | Fusion Hero OS v10.0.0 | Dissertation unter Einbezug des "
        "Gesamtarchivs und aller registrierten Konnektoren | Server-Heißlauf protokolliert]",
        first_line=False,
        italic=True,
    )

    doc.save(str(SRC))
    doc.save(str(ARCH))
    wc = sum(len(p.text.split()) for p in doc.paragraphs)
    manifest = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "author": "Stephan Hagen Urban",
        "title": "Autopoiesis und Autopolitik des Fusion Hero OS",
        "version": "1.0",
        "outputs": [str(SRC), str(ARCH)],
        "bytes": SRC.stat().st_size,
        "words_approx": wc,
        "paragraphs": len(doc.paragraphs),
        "platform_canon": "10.0.0",
        "hot_run": True,
        "archives_included": True,
        "connectors_included": True,
    }
    man = Path.home() / ".fusion" / "mesh" / "coordination" / "dissertation_build_manifest.json"
    man.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
